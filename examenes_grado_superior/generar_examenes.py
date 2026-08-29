# -*- coding: utf-8 -*-
"""
Generador de 5 exámenes de Programación (Java) para Grado Superior (1º DAW/DAM).
Genera cada examen en DOCX (python-docx) y PDF (fpdf2 con fuente Unicode).

Estructura de cada examen:
  * Portada + tabla informativa
  * Parte 1: Test / opción múltiple y verdadero-falso
  * Parte 2: Análisis / depuración (con trazas)
  * Parte 3: Ejercicios de programación (resueltos)
  * Parte 4: Pregunta teórica
  * Plantilla de espacio para respuestas

Uso:
  python generar_examenes.py            # genera DOCX y PDF de los 5 exámenes
  python generar_examenes.py docx       # solo DOCX
  python generar_examenes.py pdf        # solo PDF
"""

import os
import sys

# ---------------------------------------------------------------------------
# CONTENIDO DE LOS 5 EXÁMENES
# ---------------------------------------------------------------------------

CURSO = "1º Grado Superior — Desarrollo de Aplicaciones Web (DAW)"
ASIGNATURA = "Programación (PRG)"
DURACION = "2 horas"

# Estructura de cada examen:
#  "titulo", "subtitulo", "temas" -> lista de temas cubiertos
#  "puntuacion" -> total de puntos
#  "info" -> filas extra de la tabla informativa (etiqueta, valor)
#  "secciones" -> lista de secciones:
#     {"titulo": ..., "puntos": ..., "preguntas": [pregunta,...]}
#  Cada pregunta tiene "tipo" (test | code_analysis | ejercicio | teorica | vf)
#  y campos según tipo.
EXAMENES = []

# =====================================================================
# EXAMEN 1 — Fundamentos de Java y control de flujo
# =====================================================================
EXAMENES.append({
    "titulo": "EXAMEN DE PROGRAMACIÓN – 1º GRADO SUPERIOR",
    "subtitulo": "Examen 1 · Fundamentos de Java y control de flujo",
    "temas": [
        "Introducción a Java: JVM, bytecode, tipos primitivos y referencias",
        "Variables, operadores, conversiones y promociones",
        "Estructuras de control: if / else-if / switch, bucles while, do-while, for",
        "Entrada/salida por consola (Scanner)",
        "Métodos: paso de parámetros por valor y por referencia",
    ],
    "puntuacion": 16,
    "info": [("Convocatoria", "Evaluación 1"), ("Valoración", "Teoría y práctica sobre Java básico")],
    "secciones": [
        {
            "titulo": "Parte 1: Opción múltiple y Verdadero/Falso",
            "puntos": "5 puntos — 0,5 cada pregunta",
            "preguntas": [
                {
                    "tipo": "test",
                    "enunciado": "Dado el código siguiente, ¿qué se imprime?",
                    "code": "int a = 7;\ndouble b = 2.0;\nSystem.out.println(a / b);",
                    "opciones": ["a) 3", "b) 3.0", "c) 3.5", "d) Error de compilación"],
                    "respuesta": "c) 3.5 — Al dividir un int por un double, Java promociona el int a double y el resultado es 3.5.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Cuál de estos bucles garantiza que el cuerpo se ejecute al menos una vez?",
                    "opciones": ["a) while (cond) { }", "b) for (int i=0; i<n; i++) { }",
                                  "c) do { } while (cond);", "d) Ninguno lo garantiza"],
                    "respuesta": "c) do-while — Evalúa la condición al final, por lo que ejecuta el cuerpo al menos una vez.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Si ejecutamos números con tipo byte, ¿qué operación causaría un error de compilación por posible pérdida de precisión?",
                    "opciones": ["a) byte r = (byte) (a + b);", "b) byte r = a + b;",
                                  "c) int r = a + b;", "d) long r = a + b;"],
                    "respuesta": "b) byte r = a + b; — a + b se promociona a int al operar, por lo que asignarlo a byte requiere un cast explícito.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "En Java, cuando se pasa un array a un método, las modificaciones que haga el método sobre los elementos se reflejan en el array original.",
                    "respuesta": "Verdadero — En Java todo se pasa por valor, pero las referencias a objetos (incluidos los arrays) copian la dirección. Al modificar elementos a través de esa referencia, se modifica el objeto original.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Dado el switch, ¿qué valor se imprime si op = 2?",
                    "code": "int op = 2;\nswitch (op) {\n  case 1: System.out.print(\"A\"); break;\n  case 2: System.out.print(\"B\");\n  case 3: System.out.print(\"C\"); break;\n  default: System.out.print(\"D\");\n}",
                    "opciones": ["a) B", "b) BC", "c) BCD", "d) A"],
                    "respuesta": "b) BC — El case 2 no tiene break, por lo que 'cae' (fall-through) hacia el case 3 y también imprime C antes de llegar al break.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "En Java, todas las variables de tipo primitivo se almacenan en el montículo (heap) cuando se declaran dentro de un método.",
                    "respuesta": "Falso — Las variables locales de tipo primitivo se almacenan en la pila (stack), no en el heap.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué imprime el siguiente bucle?",
                    "code": "int i = 0;\nwhile (i < 5) {\n  if (i % 2 == 0) System.out.print(i);\n  i++;\n}",
                    "opciones": ["a) 024", "b) 135", "c) 01234", "d) 1234"],
                    "respuesta": "a) 024 — Imprime los valores de i que son pares (0, 2, 4).",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué método de Scanner permite leer un número entero correctamente?",
                    "opciones": ["a) sc.next()", "b) sc.nextInt()", "c) sc.nextDouble()", "d) sc.nextLine()"],
                    "respuesta": "b) sc.nextInt() — Lee y devuelve el siguiente token como int.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Dado el código:",
                    "code": "int x = 10;\nint z = (x > 5) ? x * 2 : x / 2;\nSystem.out.println(z);",
                    "opciones": ["a) 10", "b) 5", "c) 20", "d) 2"],
                    "respuesta": "c) 20 — El operador ternario evalúa la condición (true), por lo que z = 10 * 2 = 20.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Cuál de estas declaraciones de un método es el correcto para un método que no devuelve nada y no recibe parámetros?",
                    "opciones": ["a) public static int main()", "b) public static void ejemplo()",
                                  "c) public static String ejemplo()", "d) public static example()"],
                    "respuesta": "b) public static void ejemplo() — void indica que no devuelve nada.",
                },
            ],
        },
        {
            "titulo": "Parte 2: Análisis y depuración",
            "puntos": "3 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "El siguiente programa tiene 3 errores. Encuéntralos y explica por qué fallan.",
                    "puntos": "1,5 ptos",
                    "code": "public class Promedio {\n    public static void main(String[] args) {\n        int[] notas = {4, 7, 3, 9};\n        double media = calcular(notas[]);\n        System.out.println(\"Media: \" + media);\n    }\n    static double calcular(int[] n) {\n        int suma = 0;\n        for (int i = 0; i <= n.length; i++) {\n            suma += n[i];\n        }\n        return suma;\n    }\n}",
                    "respuesta": [
                        "1. `calcular(notas[])` → no se pasan los corchetes al invocar: debe ser `calcular(notas)`.",
                        "2. `i <= n.length` → desbordamiento: el último índice válido es length-1. Debe ser `i < n.length`.",
                        "3. `return suma;` → no devuelve la media: debe ser `return suma / (double) n.length;` para que el resultado sea 7,25 (doble división real y no entera).",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Analiza el siguiente código y determina qué imprime (traza manual).",
                    "puntos": "1,5 ptos",
                    "code": "public class Traza {\n    public static void main(String[] args) {\n        int a = 3;\n        int b = 5;\n        int c = metodo(a, b);\n        System.out.println(a + \" \" + b + \" \" + c);\n    }\n    static int metodo(int x, int y) {\n        x = x + 2;\n        y = y + 1;\n        return x * y;\n    }\n}",
                    "respuesta": [
                        "En Java los primitivos se pasan por valor: dentro de `metodo`, x se convierte en 5 e y en 6, y devuelve 5*6 = 30.",
                        "Sin embargo, a y b en `main` NO cambian porque el paso es por valor (se copian).",
                        "Salida: `3 5 30`",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 3: Ejercicios de programación",
            "puntos": "7 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "Escribe un programa que pida un número entero positivo n y calcule la suma de todos los múltiplos de 3 o de 5 menores que n.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.util.Scanner;",
                        "public class Multiplos {",
                        "    public static void main(String[] args) {",
                        "        Scanner sc = new Scanner(System.in);",
                        "        System.out.print(\"n: \");",
                        "        int n = sc.nextInt();",
                        "        int suma = 0;",
                        "        for (int i = 1; i < n; i++)",
                        "            if (i % 3 == 0 || i % 5 == 0)",
                        "                suma += i;",
                        "        System.out.println(\"Suma: \" + suma);",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Implementa un método que reciba un número entero y devuelva true si es primo y false en caso contrario. Úsalo en un programa que muestre los primeros 10 números primos.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "public class Primos {",
                        "    static boolean esPrimo(int n) {",
                        "        if (n < 2) return false;",
                        "        for (int i = 2; i <= Math.sqrt(n); i++)",
                        "            if (n % i == 0) return false;",
                        "        return true;",
                        "    }",
                        "    public static void main(String[] args) {",
                        "        int encontrados = 0, n = 2;",
                        "        while (encontrados < 10) {",
                        "            if (esPrimo(n)) { System.out.println(n); encontrados++; }",
                        "            n++;",
                        "        }",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Escribe un programa que lea dos números por teclado y muestre un menú (sumar, restar, multiplicar, dividir). Usa un bucle do-while que se repita mientras el usuario no elija salir, y gestiona la división entre cero.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.util.Scanner;",
                        "public class Calculadora {",
                        "    public static void main(String[] args) {",
                        "        Scanner sc = new Scanner(System.in);",
                        "        int opcion;",
                        "        do {",
                        "            System.out.println(\"1.Sumar 2.Restar 3.Multiplicar 4.Dividir 0.Salir\");",
                        "            opcion = sc.nextInt();",
                        "            if (opcion == 0) break;",
                        "            System.out.print(\"a y b: \");",
                        "            double a = sc.nextDouble(), b = sc.nextDouble();",
                        "            switch (opcion) {",
                        "                case 1: System.out.println(\"Suma: \" + (a + b)); break;",
                        "                case 2: System.out.println(\"Resta: \" + (a - b)); break;",
                        "                case 3: System.out.println(\"Prod: \" + (a * b)); break;",
                        "                case 4: if (b == 0) System.out.println(\"División entre cero\");",
                        "                         else System.out.println(\"Div: \" + (a / b)); break;",
                        "                default: System.out.println(\"Opción no válida\");",
                        "            }",
                        "        } while (opcion != 0);",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Explica con tus palabras la diferencia entre pasar un parámetro por valor y por referencia en Java, y entre tipos primitivos y de referencia. Pon un ejemplo en el que se aprecie la diferencia.",
                    "puntos": "1 pto",
                    "respuesta": [
                        "Java SIEMPRE pasa parámetros por valor: se copia el valor de la variable o la referencia del objeto. Para primitivos, las modificaciones dentro del método no afectan a la variable original. Para objetos/arrays, se copia la referencia, de modo que modificar el contenido (p.ej. notas[i]=5) sí afecta al original, pero reasignar la variable (`obj = new ...`) no lo afecta fuera.",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 4: Pregunta teórica",
            "puntos": "1 punto",
            "preguntas": [
                {
                    "tipo": "teorica",
                    "enunciado": "Explica la diferencia entre los bucles `while`, `do-while` y `for`. ¿Cuándo conviene usar cada uno? Razonar la respuesta.",
                    "respuesta": [
                        "`while`: evalúa la condición al inicio → 0 o más iteraciones. Útil cuando no sabemos cuántas veces iteraremos y solo conocemos la condición.",
                        "`do-while`: evalúa al final → 1 o más iteraciones. Útil cuando el cuerpo debe ejecutarse como mínimo una vez (p. ej., menús).",
                        "`for`: compacto para iteraciones con contador conocido (índice + condición + incremento). Útil para recorrer arrays o rangos determinados.",
                    ],
                },
            ],
        },
    ],
})

# =====================================================================
# EXAMEN 2 — Arrays y Strings
# =====================================================================
EXAMENES.append({
    "titulo": "EXAMEN DE PROGRAMACIÓN – 1º GRADO SUPERIOR",
    "subtitulo": "Examen 2 · Arrays y Strings",
    "temas": [
        "Arrays unidimensionales: declaración, inicialización y recorrido",
        "Arrays bidimensionales (matrices)",
        "Ordenación y búsqueda (burbuja, selección, búsqueda binaria)",
        "La clase String: inmutabilidad, métodos, comparación",
        "Ejercicios con Scanner y procesamiento de texto",
    ],
    "puntuacion": 16,
    "info": [("Convocatoria", "Evaluación 1"), ("Valoración", "Arrays, matrices y Strings")],
    "secciones": [
        {
            "titulo": "Parte 1: Opción múltiple y Verdadero/Falso",
            "puntos": "5 puntos — 0,5 cada pregunta",
            "preguntas": [
                {
                    "tipo": "test",
                    "enunciado": "Declaración correcta de un array de 5 enteros:",
                    "opciones": ["a) int [] nums = int[5];", "b) int[] nums = new int[5];",
                                  "c) int nums[]; nums = [5];", "d) int nums = new int[5];"],
                    "respuesta": "b) int[] nums = new int[5]; — Crea un array de 5 enteros, inicializados a 0.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Cuántos elementos hay en la matriz `int[][] m = {{1,2,3},{4,5}};`?",
                    "opciones": ["a) 6", "b) 5", "c) 3", "d) 2"],
                    "respuesta": "b) 5 — La primera fila tiene 3 elementos y la segunda 2: 3 + 2 = 5.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Dado:",
                    "code": "String a = \"Hola\";\nString b = new String(\"Hola\");\nSystem.out.println(a == b);",
                    "opciones": ["a) true", "b) false", "c) Error de compilación", "d) Hola"],
                    "respuesta": "b) false — `==` compara referencias. `new String(...)` crea un nuevo objeto en el heap, distinto del literal del pool, por lo que las referencias son distintas aunque el contenido sea igual.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué devuelve `\"programacion\".substring(3, 7)`?",
                    "opciones": ["a) gram", "b) ogra", "c) grama", "d) ogram"],
                    "respuesta": "a) gram — substring(3,7) devuelve los caracteres desde el índice 3 (inclusive) hasta el 7 (exclusivo): g,r,a,m = 'gram'.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "Las cadenas String en Java son mutables, es decir, su contenido se puede cambiar una vez creadas.",
                    "respuesta": "Falso — La clase String es inmutable. Cada operación que parece 'modificarla' crea una nueva cadena.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Tras ejecutar `int[] a = {5,3,8,1}; Arrays.sort(a);` el array queda:",
                    "opciones": ["a) {5,3,8,1}", "b) {1,3,5,8}", "c) {8,5,3,1}", "d) No se puede ordenar"],
                    "respuesta": "b) {1,3,5,8} — Arrays.sort ordena ascendentemente.",
                },
                {
                    "tipo": "test",
                    "enunciado": "La búsqueda binaria (binarySearch) requiere que el array esté:",
                    "opciones": ["a) Al revés", "b) Sin importar el orden", "c) Ordenado ascendentemente", "d) Lleno de números únicos"],
                    "respuesta": "c) Ordenado ascendentemente — La búsqueda binaria divide el intervalo, lo que solo es correcto si el array está ordenado.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Cuál es la sintaxis correcta para recorrer un array con for-each?",
                    "opciones": ["a) for (int i : numeros) { }", "b) for (int numeros[0]) { }",
                                  "c) for-each (int n) { }", "d) for (int i = numeros) { }"],
                    "respuesta": "a) for (int i : numeros) { } — El for-each declara una variable y usa ':' seguido del array.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué imprime este código?",
                    "code": "String s = \"Abc\";\nSystem.out.println(s.concat(\"DEF\") + s);",
                    "opciones": ["a) AbcDEFAbc", "b) AbcDEF", "c) AbcAbcDEF", "d) abcDEF"],
                    "respuesta": "a) AbcDEFAbc — concat devuelve una nueva cadena \"AbcDEF\"; como String es inmutable, s sigue siendo \"Abc\". Al concatenar queda \"AbcDEFAbc\".",
                },
                {
                    "tipo": "test",
                    "enunciado": "Para comparar el contenido de dos Strings correctamente se usa:",
                    "opciones": ["a) a == b", "b) a.equals(b)", "c) a.compare(b)", "d) String.equals(a,b)"],
                    "respuesta": "b) a.equals(b) — Compara el contenido textual; `==` compara referencias.",
                },
            ],
        },
        {
            "titulo": "Parte 2: Análisis y depuración",
            "puntos": "3 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "El siguiente código tiene 3 errores. Localízalos y explica cómo corregirlos. Además, indica qué se esperaría imprimir si funcionara.",
                    "puntos": "1,5 ptos",
                    "code": "public class ArraysBug {\n    public static void main(String[] args) {\n        int[] nums = {3, 1, 4, 1, 5};\n        System.out.println(suma(nums));\n        System.out.println(nums);\n    }\n    static int suma(int[] n) {\n        int t = 0;\n        for (int i = 0; i <= n.length; i++)\n            t += n[i];\n        return t;\n    }\n}",
                    "respuesta": [
                        "1. `i <= n.length` → fuera de límite; debe ser `i < n.length`.",
                        "2. `System.out.println(nums)` → imprime la referencia [I@..., no el contenido. Debe ser `Arrays.toString(nums)`.",
                        "3. (real) El método `suma` devolvería la suma correcta si se corrige el bucle. Con los datos, la suma esperada es 14.",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Analiza y determina la salida exacta de este código (traza).",
                    "puntos": "1,5 ptos",
                    "code": "public class MatrizT {\n    public static void main(String[] args) {\n        int[][] m = {{1, 2}, {3, 4}, {5, 6}};\n        System.out.println(\"F: \" + m.length);\n        System.out.println(\"C: \" + m[0].length);\n        System.out.println(m[2][1] - m[0][0]);\n        for (int i = 0; i < m.length; i++)\n            System.out.print(m[i][i%2] + \" \");\n    }\n}",
                    "respuesta": [
                        "m es una matriz de 3 filas y 2 columnas.",
                        "Salida: F: 3, C: 2, luego `m[2][1] - m[0][0] = 6 - 1 = 5`.",
                        "El bucle imprime m[0][0]=1, m[1][1]=4, m[2][0]=5 → `1 4 5`.",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 3: Ejercicios de programación",
            "puntos": "7 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "Pide 8 números por teclado y guárdalos en un array. Implementa métodos para: (1) contar cuántos son pares, (2) calcular el mayor y el menor, y (3) invertir el array. Muestra el resultado.",
                    "puntos": "2,5 ptos",
                    "respuesta": [
                        "import java.util.Scanner;",
                        "public class ArraysUtils {",
                        "    public static void main(String[] args) {",
                        "        Scanner sc = new Scanner(System.in);",
                        "        int[] nums = new int[8];",
                        "        for (int i = 0; i < nums.length; i++) stats: nums[i] = sc.nextInt(); // simplificado",
                        "        System.out.println(\"Pares: \" + contarPares(nums));",
                        "        System.out.println(\"Mayor: \" + mayor(nums) + \" Menor: \" + menor(nums));",
                        "        int[] inv = invertir(nums);",
                        "        System.out.println(\"Invertido: \" + java.util.Arrays.toString(inv));",
                        "    }",
                        "    static int contarPares(int[] a) { int c=0; for (int v:a) if (v%2==0) c++; return c; }",
                        "    static int mayor(int[] a) { int m=a[0]; for (int v:a) if (v>m) m=v; return m; }",
                        "    static int menor(int[] a) { int m=a[0]; for (int v:a) if (v<m) m=v; return m; }",
                        "    static int[] invertir(int[] a) {",
                        "        int[] r = new int[a.length];",
                        "        for (int i = 0; i < a.length; i++) r[i] = a[a.length-1-i];",
                        "        return r;",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Una matriz de 3x3 representa una tabla de números. Escribe un programa que la rellene con los valores de la multiplicación i*j (tabla de multiplicar) y que luego sume la diagonal principal y la diagonal secundaria.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "public class MatrizDiag {",
                        "    public static void main(String[] args) {",
                        "        int[][] m = new int[3][3];",
                        "        int diagP = 0, diagS = 0;",
                        "        for (int i = 0; i < 3; i++)",
                        "            for (int j = 0; j < 3; j++) {",
                        "                m[i][j] = i * j;",
                        "                if (i == j) diagP += m[i][j];",
                        "                if (i + j == 2) diagS += m[i][j];",
                        "            }",
                        "        System.out.println(\"Diagonal principal: \" + diagP);",
                        "        System.out.println(\"Diagonal secundaria: \" + diagS);",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Escribe un programa que, dada una cadena introducida por teclado, cuente cuántas vocales contiene, cuántas consonantes y cuántas palabras tiene (separadas por espacios).",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.util.Scanner;",
                        "public class ContadorTexto {",
                        "    public static void main(String[] args) {",
                        "        Scanner sc = new Scanner(System.in);",
                        "        String t = sc.nextLine().toLowerCase();",
                        "        int voc = 0, cons = 0;",
                        "        for (char c : t.toCharArray()) {",
                        "            if (c >= 'a' && c <= 'z') {",
                        "                if (\"aeiou\".indexOf(c) != -1) voc++; else cons++;",
                        "            }",
                        "        }",
                        "        String[] pal = t.trim().split(\"\\\\s+\");",
                        "        int n = (t.trim().isEmpty()) ? 0 : pal.length;",
                        "        System.out.println(\"Vocales: \" + voc + \" Consonantes: \" + cons + \" Palabras: \" + n);",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Escribe un programa que ordene un array usando el método de la burbuja y, tras ordenarlo, busque un valor introducido por teclado mediante búsqueda binaria, mostrando su posición o -1 si no existe.",
                    "puntos": "2,5 ptos",
                    "respuesta": [
                        "import java.util.Scanner;",
                        "import java.util.Arrays;",
                        "public class OrdenaBusca {",
                        "    static void burbuja(int[] a) {",
                        "        for (int i = 0; i < a.length-1; i++)",
                        "            for (int j = 0; j < a.length-1-i; j++)",
                        "                if (a[j] > a[j+1]) { int t=a[j]; a[j]=a[j+1]; a[j+1]=t; }",
                        "    }",
                        "    static int binaria(int[] a, int valor) {",
                        "        int izq=0, der=a.length-1;",
                        "        while (izq <= der) {",
                        "            int med = (izq+der)/2;",
                        "            if (a[med]==valor) return med;",
                        "            if (a[med] < valor) izq = med+1; else der = med-1;",
                        "        }",
                        "        return -1;",
                        "    }",
                        "    public static void main(String[] args) {",
                        "        int[] arr = {9, 3, 7, 1, 5, 2, 8, 4};",
                        "        burbuja(arr);",
                        "        System.out.println(Arrays.toString(arr));",
                        "        Scanner sc = new Scanner(System.in);",
                        "        System.out.print(\"Buscar: \");",
                        "        System.out.println(\"Posición: \" + binaria(arr, sc.nextInt()));",
                        "    }",
                        "}",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 4: Pregunta teórica",
            "puntos": "1 punto",
            "preguntas": [
                {
                    "tipo": "teorica",
                    "enunciado": "¿Por qué es importante que un array esté ordenado para aplicar búsqueda binaria? Explica el funcionamiento paso a paso y su complejidad en notación O grande.",
                    "respuesta": [
                        "La búsqueda binaria divide repetidamente el intervalo de búsqueda a la mitad comparando el elemento central. Esto solo funciona si el array está ordenado, porque permite descartar la mitad izquierda o derecha con una sola comparación.",
                        "Complejidad: O(log n) en el peor caso porque se reduce a la mitad en cada paso (log2 n pasos). Es mucho más eficiente que la búsqueda lineal O(n) para datos grandes.",
                    ],
                },
            ],
        },
    ],
})

# =====================================================================
# EXAMEN 3 — Programación Orientada a Objetos
# =====================================================================
EXAMENES.append({
    "titulo": "EXAMEN DE PROGRAMACIÓN – 1º GRADO SUPERIOR",
    "subtitulo": "Examen 3 · Programación Orientada a Objetos",
    "temas": [
        "Clases y objetos. Constructores",
        "Encapsulación: modificadores de acceso, getters y setters",
        "Atributos estáticos y métodos estáticos",
        "Sobrecarga de métodos y constructores",
        "La palabra reservada this. Visibilidad y alcance",
        "Estructuras de datos con objetos (ArrayList)",
    ],
    "puntuacion": 16,
    "info": [("Convocatoria", "Evaluación 2"), ("Valoración", "POO: clases, objetos y encapsulación")],
    "secciones": [
        {
            "titulo": "Parte 1: Opción múltiple y Verdadero/Falso",
            "puntos": "5 puntos — 0,5 cada pregunta",
            "preguntas": [
                {
                    "tipo": "test",
                    "enunciado": "¿Qué palabra clave se usa para impedir que una clase sea heredada?",
                    "opciones": ["a) static", "b) final", "c) abstract", "d) private"],
                    "respuesta": "b) final — Una clase marcada como final no puede ser heredada.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Un atributo declarado como `protected` es accesible desde:",
                    "opciones": ["a) Solo la misma clase", "b) Cualquier clase del programa",
                                  "c) La misma clase, paquete y subclases", "d) Solo la clase padre"],
                    "respuesta": "c) protected — accesible desde la misma clase, el mismo paquete y las subclases (incluso en otros paquetes).",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Cuál es el valor de `contador` tras ejecutar este código?",
                    "code": "class A {\n  static int contador = 0;\n  A() { contador++; }\n}\npublic class Main { public static void main(String[] a) {\n  new A(); new A(); new A();\n  A.contador = A.contador + 1;\n  System.out.println(A.contador);\n} }",
                    "opciones": ["a) 3", "b) 4", "c) 2", "d) 1"],
                    "respuesta": "b) 4 — Cada `new A()` incrementa el atributo estático (3), y luego se suma 1 → 4.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "Un campo `private` de una clase puede leerse desde otra clase del mismo paquete sin usar getters.",
                    "respuesta": "Falso — `private` solo es accesible desde la propia clase; las demás clases, incluso del mismo paquete, requieren accesores (getter/setter).",
                },
                {
                    "tipo": "test",
                    "enunciado": "La sobrecarga de métodos se define como:",
                    "opciones": ["a) Métodos que se llaman igual pero con distinta firma", "b) Métodos que se heredan",
                                  "c) Métodos con distinto nombre e igual parámetros", "d) Métodos que devuelven el mismo tipo"],
                    "respuesta": "a) Sobrecarga = mismo nombre, distinta lista de parámetros (tipo/número). No depende solo del tipo de retorno.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Cuál es la función de `this` dentro de un constructor?",
                    "opciones": ["a) Referirse al objeto actual", "b) Llamar al constructor padre",
                                  "c) Devolver el valor actual de la variable", "d) Crear un nuevo objeto"],
                    "respuesta": "a) this refiere al objeto actual y permite distinguir atributos de parámetros del mismo nombre.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué imprime este código?",
                    "code": "class Persona {\n  String nombre;\n  Persona(String nombre) { this.nombre = nombre; }\n}\npublic class Main {\n  public static void main(String[] a) {\n    Persona p1 = new Persona(\"Ana\");\n    Persona p2 = p1;\n    p2.nombre = \"Luis\";\n    System.out.println(p1.nombre);\n  }\n}",
                    "opciones": ["a) Ana", "b) Luis", "c) null", "d) Error"],
                    "respuesta": "b) Luis — p2 apunta al mismo objeto que p1 (se copia la referencia), por lo que `p2.nombre = \"Luis\"` cambia el nombre del objeto compartido, y p1.nombre es \"Luis\".",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué constructor se invoca con `new Animal()` si la clase Animal define un constructor con parámetros?",
                    "opciones": ["a) El constructor con parámetros por defecto", "b) Error de compilación si no existe constructor sin parámetros",
                                  "c) El constructor por defecto implícito", "d) Ninguno"],
                    "respuesta": "b) Si defines un constructor con parámetros, Java NO genera el constructor sin parámetros automáticamente, por lo que `new Animal()` provoca error de compilación a menos que lo definas.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "Un método `static` puede acceder directamente a atributos de instancia de la clase.",
                    "respuesta": "Falso — Un método estático pertenece a la clase y no puede acceder a atributos de instancia sin un objeto (no existe 'this' en contexto estático).",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Cuál es la diferencia entre `public`, `private` y `protected`?",
                    "opciones": ["a) Solo afectan a la velocidad del programa",
                                  "b) public accesible desde cualquier clase; private solo desde la propia; protected desde paquete y subclases",
                                  "c) Son sinónimos en Java",
                                  "d) No existen estos modificadores en Java"],
                    "respuesta": "b) public (todos), private (solo la clase), protected (paquete + subclases).",
                },
            ],
        },
        {
            "titulo": "Parte 2: Análisis y depuración",
            "puntos": "3 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "El siguiente código tiene errores de diseño y de sintaxis. Encuéntralos y corrige la clase para que sea correcta y encapsulada.",
                    "puntos": "1,5 ptos",
                    "code": "public class Cuenta {\n    String titular;\n    double saldo;\n    public Cuenta(String titular, double saldo) {\n        titular = titular;\n        saldo = saldo;\n    }\n    public void depositar(double m) { saldo += m; }\n    public void retirar(double m) { if (saldo >= m) saldo -= m; }\n    public void mostrar() { System.out.println(titular + \" \" + saldo); }\n}",
                    "respuesta": [
                        "1. Faltan los modificadores `private` en titular y saldo → debe ser `private String titular; private double saldo;`.",
                        "2. En el constructor, `titular = titular;` no hace nada (asigna el parámetro a sí mismo). Debe ser `this.titular = titular; this.saldo = saldo;`.",
                        "3. Para una buena encapsulación, faltan getters/setters. Se recomienda añadir `getTitular()`, `getSaldo()`, etc., y validar en `retirar`.",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Determina la salida del siguiente programa.",
                    "puntos": "1,5 ptos",
                    "code": "class Contador {\n    static int total = 0;\n    int valor;\n    Contador() { total++; valor = total; }\n    void mostrar() { System.out.println(\"v=\" + valor + \" total=\" + total); }\n}\npublic class Main {\n    public static void main(String[] a) {\n        Contador c1 = new Contador();\n        Contador c2 = new Contador();\n        Contador c3 = new Contador();\n        c1.mostrar(); c2.mostrar(); c3.mostrar();\n    }\n}",
                    "respuesta": [
                        "total es estático (compartido): 1, luego 2, luego 3.",
                        "valor es de instancia: c1.valor=1, c2.valor=2, c3.valor=3.",
                        "Salida: `v=1 total=3`, `v=2 total=3`, `v=3 total=3`.",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 3: Ejercicios de programación",
            "puntos": "7 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "Diseña la clase `Rectángulo` con atributos privados `ancho` y `alto` (double). Incluye: constructor, getters/setters, métodos `area()`, `perimetro()`, y `toString()`. Además, sobrecarga el constructor para permitir crear un cuadrado con un solo parámetro. Escribe un programa de prueba.",
                    "puntos": "2,5 ptos",
                    "respuesta": [
                        "public class Rectangulo {",
                        "    private double ancho, alto;",
                        "    public Rectangulo(double ancho, double alto) { this.ancho=ancho; this.alto=alto; }",
                        "    public Rectangulo(double lado) { this(lado, lado); }  // cuadrado",
                        "    public double getAncho() { return ancho; }",
                        "    public void setAncho(double ancho) { this.ancho = ancho; }",
                        "    public double getAlto() { return alto; }",
                        "    public void setAlto(double alto) { this.alto = alto; }",
                        "    public double area() { return ancho * alto; }",
                        "    public double perimetro() { return 2*(ancho+alto); }",
                        "    public String toString() { return ancho + \"x\" + alto; }",
                        "}",
                        "// Main:",
                        "public class Main { public static void main(String[] a) {",
                        "    Rectangulo r = new Rectangulo(4, 6);",
                        "    Rectangulo c = new Rectangulo(5);",
                        "    System.out.println(r + \" área \" + r.area());",
                        "    System.out.println(c + \" perímetro \" + c.perimetro());",
                        "} }",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Crea la clase `Estudiante` con atributos `nombre`, `edad` y una lista de notas (ArrayList<Double>). Añade métodos para añadir nota, calcular la media, determinar si está aprobado (media >= 5) y mostrar sus datos.",
                    "puntos": "2,5 ptos",
                    "respuesta": [
                        "import java.util.ArrayList;",
                        "public class Estudiante {",
                        "    private String nombre;",
                        "    private int edad;",
                        "    private ArrayList<Double> notas = new ArrayList<>();",
                        "    public Estudiante(String nombre, int edad) { this.nombre=nombre; this.edad=edad; }",
                        "    public void addNota(double n) { notas.add(n); }",
                        "    public double media() {\n       double s=0; for (double n: notas) s+=n; return s/notas.size(); }",
                        "    public boolean aprobado() { return media() >= 5; }",
                        "    public String toString() { return nombre + \" (\" + edad + \") media=\" + media(); }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Crea una clase `CuentaBancaria` con titular, saldo y número de cuenta. Los métodos retirar y depositar deben actualizar el saldo. Un atributo estático `totalCuentas` debe contar cuántas cuentas se han creado. Escribe un programa que cree varias cuentas, realice depósitos y retiradas, y muestre el total de cuentas.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "public class CuentaBancaria {",
                        "    private String titular;",
                        "    private double saldo;",
                        "    private static int totalCuentas = 0;",
                        "    public CuentaBancaria(String titular, double saldo) { this.titular=titular; this.saldo=saldo; totalCuentas++; }",
                        "    public void depositar(double m) { saldo += m; }",
                        "    public boolean retirar(double m) {",
                        "        if (saldo >= m) { saldo -= m; return true; }",
                        "        return false;",
                        "    }",
                        "    public static int getTotalCuentas() { return totalCuentas; }",
                        "    public String toString() { return titular + \": \" + saldo; }",
                        "}",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 4: Pregunta teórica",
            "puntos": "1 punto",
            "preguntas": [
                {
                    "tipo": "teorica",
                    "enunciado": "Explica el concepto de encapsulación. ¿Por qué es importante declarar los atributos como `private` y proporcionar getters/setters? ¿Qué beneficios aporta al mantenimiento del código?",
                    "respuesta": [
                        "La encapsulación es el principio que oculta el estado interno de un objeto y solo expone una interfaz pública controlada.",
                        "Declarar atributos private evita que la clase se modifique desde fuera sin control, permitiendo validar los valores (p. ej., impedir saldos negativos) y cambiando la implementación interna sin romper a los clientes de la clase.",
                        "Beneficios: seguridad, control de acceso, mantenibilidad, y facilita el testing y la evolución del código.",
                    ],
                },
            ],
        },
    ],
})

# =====================================================================
# EXAMEN 4 — Herencia, polimorfismo, interfaces y clases abstractas
# =====================================================================
EXAMENES.append({
    "titulo": "EXAMEN DE PROGRAMACIÓN – 1º GRADO SUPERIOR",
    "subtitulo": "Examen 4 · Herencia, polimorfismo, interfaces y clases abstractas",
    "temas": [
        "Herencia (extends), super, override",
        "Polimorfismo y enlace dinámico",
        "Clases abstractas y métodos abstractos",
        "Interfaces (implements)",
        "Modificadores final y static en herencia",
        "Uso de ArrayList con tipos polimórficos",
    ],
    "puntuacion": 16,
    "info": [("Convocatoria", "Evaluación 2"), ("Valoración", "Herencia y polimorfismo")],
    "secciones": [
        {
            "titulo": "Parte 1: Opción múltiple y Verdadero/Falso",
            "puntos": "5 puntos — 0,5 cada pregunta",
            "preguntas": [
                {
                    "tipo": "test",
                    "enunciado": "Para hacer que la clase `Perro` herede de `Animal` se usa:",
                    "opciones": ["a) class Perro implements Animal", "b) class Perro extends Animal",
                                  "c) class Perro inherits Animal", "d) class Perro : Animal"],
                    "respuesta": "b) class Perro extends Animal — extends indica herencia de clase; implements se usa con interfaces.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué palabra clave se usa para llamar al constructor de la clase padre dentro de un constructor hijo?",
                    "opciones": ["a) this()", "b) base()", "c) super()", "d) parent()"],
                    "respuesta": "c) super() — invoca el constructor del padre. this() invoca un constructor de la propia clase.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "Una clase abstracta puede contener métodos concretos (con implementación).",
                    "respuesta": "Verdadero — Una clase abstracta puede mezclar métodos abstractos y métodos concretos.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Dado:",
                    "code": "class Animal { void sonido() { System.out.println(\"genérico\"); } }\nclass Perro extends Animal { void sonido() { System.out.println(\"guau\"); } }\nAnimal a = new Perro();\na.sonido();",
                    "opciones": ["a) genérico", "b) guau", "c) Error de compilación", "d) Nada"],
                    "respuesta": "b) guau — El enlace dinámico (polimorfismo) invoca la implementación de la clase real del objeto (Perro) aunque la referencia sea de tipo Animal.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Una interfaz en Java:",
                    "opciones": ["a) Solo puede tener métodos abstractos (hasta Java 8 también default/static)",
                                  "b) Puede ser instanciada con new", "c) Solo puede heredar de una interfaz",
                                  "d) No puede tener constantes"],
                    "respuesta": "a) Una interfaz declara métodos abstractos (y desde Java 8 métodos default/static) que las clases que la implementan deben completar. No se puede instanciar con new.\n\n💡 **¿Qué es un método abstracto?** Es un método que solo tiene su firma o cabecera (nombre, parámetros y tipo de retorno) pero **no tiene cuerpo ni código** ({ ... }); termina con punto y coma (;). La clase que implementa la interfaz está obligada a programar su funcionamiento.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Para que una clase implemente una interfaz se usa la palabra clave:",
                    "opciones": ["a) extends", "b) uses", "c) implements", "d) inherits"],
                    "respuesta": "c) implements — Una clase implementa una interfaz usando `implements`; puede implementar varias interfaces separadas por comas.\n\n```java\ninterface Volador { void volar(); }\ninterface Nadador { void nadar(); }\n\n// Una clase implementando múltiples interfaces:\nclass Pato implements Volador, Nadador {\n    public void volar() { System.out.println(\"Volando\"); }\n    public void nadar() { System.out.println(\"Nadando\"); }\n}\n```",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Cuál es el propósito de `@Override`?",
                    "opciones": ["a) Crear un método nuevo", "b) Indicar que se sobrescribe un método heredado",
                                  "c) Hacer el método estático", "d) Hacer el método público"],
                    "respuesta": "b) @Override es una anotación que indica que el método sobrescribe un método heredado; ayuda al compilador a detectar errores de firma.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "En Java una clase puede heredar de varias clases simultáneamente (herencia múltiple).",
                    "respuesta": "Falso — Java no permite herencia múltiple de clases. Sí permite implementar varias interfaces.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué elemento NO se puede heredar en una subclase?",
                    "opciones": ["a) Métodos public", "b) Atributos protected", "c) Constructores", "d) Métodos protected"],
                    "respuesta": "c) Los constructores no se heredan; la subclase debe definir sus propios constructores e invocar al del padre usando `super()`.\n\n```java\nclass Persona {\n    Persona(String nombre) { System.out.println(\"Padre: \" + nombre); }\n}\n\nclass Empleado extends Persona {\n    Empleado(String nombre) {\n        super(nombre); // Invoca el constructor del padre con super()\n    }\n}\n```",
                },
                {
                    "tipo": "test",
                    "enunciado": "Dado un método `abstract`, la clase que lo contiene debe ser:",
                    "opciones": ["a) final", "b) abstract", "c) static", "d) Una interfaz o clase abstracta"],
                    "respuesta": "b) abstract — Un método abstract declarado obliga a que su clase sea abstracta o sea una interfaz.",
                },
            ],
        },
        {
            "titulo": "Parte 2: Análisis y depuración",
            "puntos": "3 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "El siguiente código tiene errores de herencia. Encuéntralos y corrige el diseño.",
                    "puntos": "1,5 ptos",
                    "code": "abstract class Vehiculo {\n    abstract void mover();\n}\nclass Coche extends Vehiculo { }\nclass Moto extends Vehiculo {\n    void mover() { System.out.println(\"La moto avanza\"); }\n}\npublic class Main { public static void main(String[] a) {\n    Vehiculo v = new Coche();  // ¿compila?\n    v.mover();\n} }",
                    "respuesta": [
                        "1. La clase abstracta Vehiculo tiene un método abstracto `mover()`. La clase Coche no lo implementa y NO es abstracta, por lo que NO compila (debe implementar mover() o declararse abstract).",
                        "2. Corregir: `class Coche extends Vehiculo { void mover() { System.out.println(\"El coche avanza\"); } }`.",
                        "3. Moto implementa mover() correctamente; con la corrección la salida sería \"El coche avanza\".",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Determina la salida de este programa.",
                    "puntos": "1,5 ptos",
                    "code": "abstract class Figura {\n    abstract double area();\n    void mostrar() { System.out.println(\"Área: \" + area()); }\n}\nclass Circulo extends Figura {\n    private double radio;\n    Circulo(double radio) { this.radio = radio; }\n    double area() { return Math.PI * radio * radio; }\n}\nclass Cuadrado extends Figura {\n    private double lado;\n    Cuadrado(double lado) { this.lado = lado; }\n    double area() { return lado * lado; }\n}\npublic class Main { public static void main(String[] a) {\n    Figura[] f = { new Circulo(2), new Cuadrado(3) };\n    for (Figura x : f) x.mostrar();\n} }",
                    "respuesta": [
                        "Polimorfismo: cada Figura invoca su propia implementación de area().",
                        "Circulo de radio 2 → área = π*4 ≈ 12,566.",
                        "Cuadrado de lado 3 → área = 9.",
                        "Salida: `Área: 12.566370614359172` y `Área: 9.0`.",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 3: Ejercicios de programación",
            "puntos": "7 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "Crea una jerarquía de clases para `Empleado`, `Directivo` (hereda) y `Tecnico` (hereda). Empleado tiene nombre y salario. Directivo añade `departamento`; Técnico añade `especialidad`. Sobrescribe toString en las tres y demuestra el polimorfismo con un método que reciba un Empleado e imprima su información.",
                    "puntos": "2,5 ptos",
                    "respuesta": [
                        "class Empleado {",
                        "    private String nombre; private double salario;",
                        "    public Empleado(String nombre, double salario) { this.nombre=nombre; this.salario=salario; }",
                        "    public String toString() { return nombre + \", \" + salario + \" €\"; }",
                        "}",
                        "class Directivo extends Empleado {",
                        "    private String departamento;",
                        "    public Directivo(String n, double s, String dep) { super(n, s); this.departamento=dep; }",
                        "    public String toString() { return super.toString() + \" [\" + departamento + \"]\"; }",
                        "}",
                        "class Tecnico extends Empleado {",
                        "    private String especialidad;",
                        "    public Tecnico(String n, double s, String esp) { super(n, s); this.especialidad=esp; }",
                        "    public String toString() { return super.toString() + \" <\" + especialidad + \">\"; }",
                        "}",
                        "public class Main {",
                        "    static void mostrar(Empleado e) { System.out.println(e); }",
                        "    public static void main(String[] a) {",
                        "        mostrar(new Directivo(\"Ana\", 3000, \"IT\"));",
                        "        mostrar(new Tecnico(\"Luis\", 2000, \"Redes\"));",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Define una interfaz `Movible` con el método `void moverse()`. Define también una interfaz `Sonoro` con `void sonar()`. Crea una clase `Robot` que implemente ambas interfaces y una clase `Coche` que implemente solo Movible. Escribe un programa que cree objetos y los trate a través de las interfaces.",
                    "puntos": "2,5 ptos",
                    "respuesta": [
                        "interface Movible { void moverse(); }",
                        "interface Sonoro { void sonar(); }",
                        "class Robot implements Movible, Sonoro {",
                        "    public void moverse() { System.out.println(\"El robot se mueve\"); }",
                        "    public void sonar() { System.out.println(\"Bip bip\"); }",
                        "}",
                        "class Coche implements Movible {",
                        "    public void moverse() { System.out.println(\"El coche avanza\"); }",
                        "}",
                        "public class Main {",
                        "    public static void main(String[] a) {",
                        "        Movible[] m = { new Robot(), new Coche() };",
                        "        for (Movible x : m) x.moverse();",
                        "        ((Sonoro)m[0]).sonar();",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Crea una clase abstracta `Instrumento` con un método abstracto `tocar()`. Deriva al menos tres instrumentos (Piano, Guitarra, Bateria). Prepara un arreglo de Instrumento con varios objetos y haz que todos 'toquen' usando polimorfismo. Añade un contador estático de instrumentos creados.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "abstract class Instrumento {",
                        "    static int total = 0;",
                        "    Instrumento() { total++; }",
                        "    abstract void tocar();",
                        "}",
                        "class Piano extends Instrumento { void tocar() { System.out.println(\"Piano: plin plin\"); } }",
                        "class Guitarra extends Instrumento { void tocar() { System.out.println(\"Guitarra: strum\"); } }",
                        "class Bateria extends Instrumento { void tocar() { System.out.println(\"Batería: bom bom\"); } }",
                        "public class Main { public static void main(String[] a) {",
                        "    Instrumento[] orq = { new Piano(), new Guitarra(), new Bateria(), new Piano() };",
                        "    for (Instrumento i : orq) i.tocar();",
                        "    System.out.println(\"Total instrumentos: \" + Instrumento.total);",
                        "} }",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 4: Pregunta teórica",
            "puntos": "1 punto",
            "preguntas": [
                {
                    "tipo": "teorica",
                    "enunciado": "Explica la diferencia entre una clase abstracta y una interfaz en Java. Da un ejemplo de cuándo elegirías cada una.",
                    "respuesta": [
                        "Clase abstracta: puede tener atributos de instancia, constructores y métodos concretos. Solo se puede heredar de UNA. Ideal para compartir estado/implementación común entre clases muy relacionadas (p. ej., Figuras que comparten comportamiento).",
                        "Interfaz: declara un 'contrato' (métodos) que las clases implementan; desde Java 8 puede tener métodos default/static y constantes. Una clase puede implementar VARIAS. Ideal para definir capacidades ('puede volar', 'puede moverse') reutilizables entre clases no relacionadas.",
                        "Regla general: cuando hay una relación 'es-un' fuerte y código compartido → clase abstracta; cuando se quiere un contrato flexible o múltiples capacidades → interfaz.",
                    ],
                },
            ],
        },
    ],
})

# =====================================================================
# EXAMEN 5 — Ficheros, excepciones y colecciones
# =====================================================================
EXAMENES.append({
    "titulo": "EXAMEN DE PROGRAMACIÓN – 1º GRADO SUPERIOR",
    "subtitulo": "Examen 5 · Ficheros, excepciones y colecciones",
    "temas": [
        "Lectura y escritura de ficheros de texto (Scanner, File, PrintWriter)",
        "Gestión de excepciones: try/catch, finally, try-with-resources",
        "Excepciones comprobadas y no comprobadas. jerarquía Throwable",
        "Colecciones: ArrayList, HashMap, HashSet",
        "Uso combinado de ficheros y colecciones",
    ],
    "puntuacion": 16,
    "info": [("Convocatoria", "Evaluación 2"), ("Valoración", "Ficheros, excepciones y colecciones")],
    "secciones": [
        {
            "titulo": "Parte 1: Opción múltiple y Verdadero/Falso",
            "puntos": "5 puntos — 0,5 cada pregunta",
            "preguntas": [
                {
                    "tipo": "test",
                    "enunciado": "¿Qué bloque se ejecuta SIEMPRE, haya o no excepción?",
                    "opciones": ["a) try", "b) catch", "c) finally", "d) throws"],
                    "respuesta": "c) finally — El bloque finally se ejecuta en todos los casos (con o sin excepción), salvo System.exit().",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué clase es la base de todas las excepciones en Java?",
                    "opciones": ["a) Error", "b) Exception", "c) Throwable", "d) RuntimeException"],
                    "respuesta": "c) Throwable — Es la raíz de la jerarquía; de él derivan Exception y Error.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Una excepción comprobada (checked) obliga a:",
                    "opciones": ["a) Capturarla o declararla con throws", "b) Ignorarla", "c) Lanzarla siempre",
                                  "d) Convertirla a RuntimeException"],
                    "respuesta": "a) Las excepciones comprobadas deben capturarse (try/catch) o declararse en la firma con throws.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "La sentencia try-with-resources cierra automáticamente los recursos definidos dentro de ella.",
                    "respuesta": "Verdadero — try-with-resources (Java 7+) cierra automáticamente los recursos AutoCloseable al final del bloque, sin necesidad de finally.",
                },
                {
                    "tipo": "test",
                    "enunciado": "La clase `PrintWriter` se usa principalmente para:",
                    "opciones": ["a) Leer de un fichero", "b) Escribir texto formateado en un fichero de salida",
                                  "c) Crear un directorio", "d) Borrar un fichero"],
                    "respuesta": "b) PrintWriter escribe texto, a menudo con formato (printf/println), en un archivo de salida.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué estructura permite almacenar pares clave-valor y no admite claves duplicadas?",
                    "opciones": ["a) ArrayList", "b) HashSet", "c) HashMap", "d) LinkedList"],
                    "respuesta": "c) HashMap — Almacena pares clave-valor; las claves son únicas (si se inserta una clave existente, se reemplaza el valor).",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué imprime este código?",
                    "code": "List<Integer> l = new ArrayList<>();\nl.add(10); l.add(20); l.add(10);\nSystem.out.println(l.size());",
                    "opciones": ["a) 2", "b) 3", "c) 4", "d) Error"],
                    "respuesta": "b) 3 — ArrayList permite elementos duplicados, por lo que hay 3 elementos pese a repetir el 10.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "Un HashSet garantiza el orden de los elementos en la inserción.",
                    "respuesta": "Falso — HashSet no garantiza ningún orden de los elementos; para orden de inserción se usa LinkedHashSet o List.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué ocurre al intentar leer un fichero que no existe?",
                    "opciones": ["a) Se crea automáticamente", "b) Se lanza FileNotFoundException",
                                  "c) Devuelve un fichero vacío", "d) Se lanza NumberFormatException"],
                    "respuesta": "b) FileNotFoundException (subclase de IOException, comprobada) se lanza cuando se intenta abrir un fichero inexistente para lectura.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Para lanzar una excepción propia dentro de un método se usa la palabra clave:",
                    "opciones": ["a) throw", "b) throws", "c) try", "d) new"],
                    "respuesta": "a) throw lanza una excepción; `throws` declara las excepciones que un método puede propagar.",
                },
            ],
        },
        {
            "titulo": "Parte 2: Análisis y depuración",
            "puntos": "3 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "El siguiente código tiene errores. Encuéntralos y corrígelos.",
                    "puntos": "1,5 ptos",
                    "code": "import java.io.*;\npublic class LeerArchivo {\n    public static void main(String[] args) {\n        try {\n            BufferedReader br = new BufferedReader(new FileReader(\"datos.txt\"));\n            String linea;\n            while ((linea = br.readLine()) != null) {\n                System.out.println(linea);\n            }\n        }\n        System.out.println(\"Fin\");\n    }\n}",
                    "respuesta": [
                        "1. No se captura la IOException, que el compilador exige (checked). Falta `catch (IOException e) { ... }`.",
                        "2. El BufferedReader no se cierra. Mejor usar try-with-resources: `try (BufferedReader br = new BufferedReader(new FileReader(\"datos.txt\")))`.",
                        "3. (opcional de diseño) Acepta también declarar `throws IOException` en main, aunque lo preferible es capturar la excepción.",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Determina la salida de este código (traza con excepciones).",
                    "puntos": "1,5 ptos",
                    "code": "public class ExpTraza {\n    static int dividir(int a, int b) {\n        try { return a / b; }\n        catch (ArithmeticException e) { System.out.println(\"Arithmetic\"); return -1; }\n        finally { System.out.println(\"finally\"); }\n    }\n    public static void main(String[] a) {\n        System.out.println(dividir(10, 2));\n        System.out.println(dividir(10, 0));\n    }\n}",
                    "respuesta": [
                        "Primera llamada (10/2): no hay excepción; se ejecuta finally (imprime 'finally') y devuelve 5 → imprime 5.",
                        "Segunda llamada (10/0): se lanza ArithmeticException, se captura (imprime 'Arithmetic'), luego finally (imprime 'finally'), devuelve -1 → imprime -1.",
                        "Salida completa: `finally 5 Arithmetic finally -1`.",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 3: Ejercicios de programación",
            "puntos": "7 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "Escribe un programa que lea un archivo `alumnos.txt` donde cada línea contiene `nombre nota1 nota2 nota3` separados por espacio. Calcula la media de cada alumno y escribe en `resultados.txt` el nombre y la media con 2 decimales. Gestiona las excepciones con try-with-resources.",
                    "puntos": "2,5 ptos",
                    "respuesta": [
                        "import java.io.*;",
                        "import java.util.Scanner;",
                        "public class GestionAlumnos {",
                        "    public static void main(String[] args) {",
                        "        try (Scanner sc = new Scanner(new File(\"alumnos.txt\"));",
                        "             PrintWriter pw = new PrintWriter(\"resultados.txt\")) {",
                        "            while (sc.hasNextLine()) {",
                        "                String linea = sc.nextLine().trim();",
                        "                if (linea.isEmpty()) continue;",
                        "                String[] p = linea.split(\" \");",
                        "                double s = 0; int n = p.length - 1;",
                        "                for (int i = 1; i < p.length; i++) s += Double.parseDouble(p[i]);",
                        "                pw.printf(\"%s: %.2f%n\", p[0], s / n);",
                        "            }",
                        "            System.out.println(\"resultados.txt generado\");",
                        "        } catch (FileNotFoundException e) {",
                        "            System.out.println(\"No se encuentra alumnos.txt\");",
                        "        } catch (Exception e) {",
                        "            System.out.println(\"Error: \" + e.getMessage());",
                        "        }",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Crea un programa que use un HashMap para guardar el número de habitantes de varias ciudades (clave=ciudad, valor=habitantes). Pide ciudades por teclado, inserta varias, y luego: recorre el mapa mostrando cada ciudad y su población, y muestra la ciudad con más habitantes.",
                    "puntos": "2,5 ptos",
                    "respuesta": [
                        "import java.util.*;",
                        "public class Ciudades {",
                        "    public static void main(String[] args) {",
                        "        Map<String, Integer> hab = new HashMap<>();",
                        "        hab.put(\"Madrid\", 3200000);",
                        "        hab.put(\"Barcelona\", 1600000);",
                        "        hab.put(\"Valencia\", 790000);",
                        "        hab.put(\"Madrid\", 3300000); // reemplaza",
                        "        String mayor = null; int max = -1;",
                        "        for (Map.Entry<String,Integer> e : hab.entrySet()) {",
                        "            System.out.println(e.getKey() + \" -> \" + e.getValue());",
                        "            if (e.getValue() > max) { max = e.getValue(); mayor = e.getKey(); }",
                        "        }",
                        "        System.out.println(\"Más poblada: \" + mayor);",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Escribe un método que valide si un número de hora es correcto (entre 0 y 23) y lance una excepción personalizada `HoraInvalidaException` si no lo es. Crea la excepción y un programa que pida una hora y la muestre, capturando la excepción.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "class HoraInvalidaException extends Exception {",
                        "    public HoraInvalidaException(String msg) { super(msg); }",
                        "}",
                        "import java.util.Scanner;",
                        "public class Hora {",
                        "    static void validar(int h) throws HoraInvalidaException {",
                        "        if (h < 0 || h > 23) throw new HoraInvalidaException(\"Hora fuera de rango: \" + h);",
                        "    }",
                        "    public static void main(String[] a) {",
                        "        Scanner sc = new Scanner(System.in);",
                        "        System.out.print(\"Hora: \");",
                        "        int h = sc.nextInt();",
                        "        try { validar(h); System.out.println(\"Hora válida: \" + h); }",
                        "        catch (HoraInvalidaException e) { System.out.println(e.getMessage()); }",
                        "    }",
                        "}",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 4: Pregunta teórica",
            "puntos": "1 punto",
            "preguntas": [
                {
                    "tipo": "teorica",
                    "enunciado": "Explica la diferencia entre ArrayList y HashMap, y entre ArrayList y HashSet. ¿Cuándo usarías cada una?",
                    "respuesta": [
                        "ArrayList: lista ordenada indexada, admite duplicados, acceso por posición. Ideal para una secuencia ordenada o cuando importa el índice/orden.",
                        "HashMap: almacena pares clave→valor, claves únicas, sin orden garantizado, acceso por clave en O(1). Ideal para búsquedas por clave (diccionarios).",
                        "HashSet: colección de elementos únicos, sin orden, no permite duplicados, operaciones de pertenencia eficientes. Ideal para eliminar duplicados o comprobar presencia.",
                        "Regla: necesitas orden/posición → ArrayList; clave→valor → HashMap; únicidad/presencia → HashSet.",
                    ],
                },
            ],
        },
    ],
})

# ---------------------------------------------------------------------------
# RENDERIZADO A DOCX
# ---------------------------------------------------------------------------

def render_docx(exam, out_path, dark=False):
    from docx import Document
    from docx.shared import Pt, Cm, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Paleta: clara u oscura (mismos colores que el PDF)
    if dark:
        PAL = dict(bg="181A21", text="E2E6F0", title="81B2FF", sub="AAAFBE",
                   muted="969BAA", code="CED4E2", code_bg="262A36", table="282C3A",
                   green="7AD090", green_bg="1A2A1E", line="5F6473", pink="FFB6C1")
        # Fondo de página oscuro
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        bg = OxmlElement("w:background")
        bg.set(qn("w:color"), PAL["bg"])
        doc.settings.element.append(bg)
    else:
        PAL = dict(bg="FFFFFF", text="000000", title="1A478A", sub="555555",
                   muted="8C8C8C", code="323232", code_bg="F5F7FA", table="EEF3FA",
                   green="1B5E20", green_bg="E8F5E9", line="B4B4B4", pink="FFB6C1")

    def C(h):
        return RGBColor.from_string(h)

    # Márgenes
    for section in doc.sections:
        # Misma página A4 que el PDF
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = C(PAL["text"])

    AZUL = C(PAL["title"])
    VERDE = C(PAL["green"])
    ROSA = C(PAL["pink"])
    MUTED = C(PAL["muted"])
    SUB = C(PAL["sub"])

    def h1(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = AZUL
        return p

    def h2(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = False
        r.font.size = Pt(12)
        r.font.color.rgb = SUB
        return p

    def section_title(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = AZUL
        # línea inferior
        p.paragraph_format.space_before = Pt(14)
        return p

    def table_borders(t, color):
        """Color de todos los bordes de una tabla."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tblPr = t._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            borders.append(el)
        tblPr.append(borders)

    def info_table(rows):
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = "Table Grid"
        t.autofit = True
        for i, (k, v) in enumerate(rows):
            c0 = t.rows[i].cells[0]
            c1 = t.rows[i].cells[1]
            c0.text = k
            c1.text = v
            for p in c0.paragraphs:
                for r in p.runs:
                    r.bold = True
            shade_cell(c0, PAL["table"])  # azul como en el PDF
        table_borders(t, PAL["line"])
        return t

    def add_code(code):
        for line in code.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line.replace("<", "<").replace(">", ">"))
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = C(PAL["code"])
            shade(p, PAL["code_bg"])  # fondo como el recuadro del PDF
        # espaciado
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)
        sp.add_run("").font.size = Pt(2)

    def add_par(text, indent=0.0, italic=False, color=None, bold_prefix=None):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
        r = p.add_run(text)
        r.italic = italic if italic else None
        if color:
            r.font.color.rgb = color
        return p

    def shade(p, fill="E8F5E9"):
        """Sombreado de fondo para un párrafo (igual que el PDF)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        pPr.append(shd)

    def shade_cell(cell, fill):
        """Fondo de una celda de tabla."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    def cell_border(cell, color, sz=8):
        """Borde completo de celda (sz en octavos de punto)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            borders.append(el)
        tcPr.append(borders)

    def cell_margins(cell, top=80, left=120, bottom=80, right=120):
        """Márgenes internos de una celda (en veinteavos de punto)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tcPr = cell._tc.get_or_add_tcPr()
        mar = OxmlElement("w:tcMar")
        for edge, w in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:w"), str(w))
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        tcPr.append(mar)

    def p_border_bottom(p, color, sz=6):
        """Línea horizontal inferior en un párrafo (separadores del PDF)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(sz))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        pbdr.append(bottom)
        pPr.append(pbdr)

    # --- Portada ---
    h1(exam["titulo"])
    h2(exam["subtitulo"])
    add_par(" ".join(exam["temas"]), indent=0.0, italic=True, color=MUTED)
    doc.add_paragraph()

    info_rows = [
        ("Asignatura", ASIGNATURA),
        ("Curso", CURSO),
        ("Duración", DURACION),
        ("Puntuación total", f'{exam["puntuacion"]} puntos'),
    ]
    info_rows += exam["info"]
    info_table(info_rows)

    # --- Secciones ---
    for sec in exam["secciones"]:
        doc.add_page_break() if sec is not exam["secciones"][0] else None
        section_title(sec["titulo"])
        add_par(sec["puntos"], indent=0.0, italic=True, color=MUTED)

        for num, q in enumerate(sec["preguntas"], start=1):
            # Enunciado: número en rosa pastel + texto en negrita (igual que el PDF)
            enun = q.get("enunciado", "")
            if "puntos" in q:
                enun = f"{enun}  ({q['puntos']})"
            p_q = doc.add_paragraph()
            p_q.paragraph_format.left_indent = Cm(0.5)
            rn = p_q.add_run(f"{num}. ")
            rn.bold = True
            rn.font.size = Pt(11)
            rn.font.color.rgb = ROSA
            re_ = p_q.add_run(enun)
            re_.bold = True
            re_.font.size = Pt(11)

            if q.get("code"):
                add_code(q["code"])

            if q.get("opciones"):
                for o in q["opciones"]:
                    add_par(o, indent=1.2)

            if q.get("respuesta"):
                for _ in range(6):
                    doc.add_paragraph()
                resp = q["respuesta"] if isinstance(q["respuesta"], list) else [q["respuesta"]]
                # Recuadro verde con borde (igual que el bloque del PDF)
                t = doc.add_table(rows=1, cols=1)
                t.autofit = True
                cell = t.rows[0].cells[0]
                shade_cell(cell, PAL["green_bg"])
                cell_border(cell, PAL["green"], sz=8)
                cell_margins(cell)
                p0 = cell.paragraphs[0]
                r = p0.add_run("RESPUESTA / SOLUCIÓN")
                r.bold = True
                r.font.color.rgb = VERDE
                for l in resp:
                    p2 = cell.add_paragraph()
                    r2 = p2.add_run(l)
                    r2.font.color.rgb = VERDE
            # Línea separadora rosa pastel (igual que el PDF)
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(14)
            sep.paragraph_format.space_after = Pt(4)
            p_border_bottom(sep, "FFB6C1", sz=6)

    doc.add_page_break()
    section_title("Plantilla de respuestas")
    add_par("(Páginas en blanco para desarrollar las soluciones)", italic=True, color=MUTED)
    for _ in range(14):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p_border_bottom(p, PAL["line"], sz=6)

    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# RENDERIZADO A PDF (fpdf2 con fuente Unicode)
# ---------------------------------------------------------------------------

FONT_DIRS = [
    r"C:\Windows\Fonts",
    r"C:\Users\Pablo\AppData\Local\Microsoft\Windows\Fonts",
]

def find_font(name):
    import glob
    for d in FONT_DIRS:
        for f in glob.glob(d + "\\" + name):
            if os.path.exists(f):
                return f
    return None


class PDFExam:
    def __init__(self, exam, dark=False):
        from fpdf import FPDF

        class _PDF(FPDF):
            """FPDF que rellena el fondo de cada página (necesario en modo oscuro)."""

            def __init__(self, bg):
                super().__init__("P", "mm", "A4")
                self.bg = bg

            def header(self):
                self.set_fill_color(*self.bg)
                self.rect(0, 0, 210, 297, "F")

        self.exam = exam
        self.dark = dark
        # Paleta de colores: clara (por defecto) u oscura
        if dark:
            self.c_bg = (24, 26, 33)
            self.c_text = (226, 230, 240)
            self.c_title = (129, 178, 255)
            self.c_sub = (170, 175, 190)
            self.c_muted = (150, 155, 170)
            self.c_code = (206, 212, 226)
            self.c_code_bg = (38, 42, 54)
            self.c_table = (40, 44, 58)
            self.c_green = (122, 208, 144)
            self.c_green_bg = (26, 42, 30)
            self.c_line = (95, 100, 115)
            self.c_pink = (255, 182, 193)
        else:
            self.c_bg = (255, 255, 255)
            self.c_text = (0, 0, 0)
            self.c_title = (26, 71, 138)
            self.c_sub = (85, 85, 85)
            self.c_muted = (140, 140, 140)
            self.c_code = (50, 50, 50)
            self.c_code_bg = (245, 247, 250)
            self.c_table = (238, 243, 250)
            self.c_green = (27, 94, 32)
            self.c_green_bg = (232, 245, 233)
            self.c_line = (180, 180, 180)
            self.c_pink = (255, 182, 193)
        self.pdf = _PDF(self.c_bg)
        self.pdf.set_auto_page_break(auto=True, margin=15)
        self.page_w = 210
        self.margin = 20
        self.content_w = self.page_w - 2 * self.margin

        # Registrar fuentes Unicode (Arial para texto y Courier para código)
        self.pdf.add_font("Ar", "", find_font("arial.ttf"), uni=True)
        self.pdf.add_font("Ar", "B", find_font("arialbd.ttf"), uni=True)
        self.pdf.add_font("Ar", "I", find_font("ariali.ttf"), uni=True)
        self.pdf.add_font("Ar", "BI", find_font("arialbi.ttf"), uni=True)
        self.pdf.add_font("Cou", "", find_font("cour.ttf"), uni=True)
        self.pdf.add_font("Cou", "B", find_font("courbd.ttf"), uni=True)

    # Utilidades
    def check_page(self, h):
        if self.pdf.get_y() + h > 287:
            self.pdf.add_page()

    def para(self, text, size=11, style="", color=None, h=5.5, indent=0.0):
        w = self.content_w - indent * 10
        self.check_page(10)
        x = self.margin + indent * 10
        self.pdf.set_xy(x, self.pdf.get_y())
        self.pdf.set_font("Ar", style, size)
        self.pdf.set_text_color(*(color if color is not None else self.c_text))
        self.pdf.multi_cell(w, h, text, align="L")
        self.pdf.ln(1.5)

    def separador(self):
        """Línea horizontal rosa pastel para separar ejercicios."""
        p = self.pdf
        y = p.get_y() + 22  # al menos 4 líneas (5.5 mm c/u) antes de la línea
        if y + 8 > 280:
            p.ln(8)
            return
        p.set_draw_color(*self.c_pink)
        p.set_line_width(0.6)
        p.line(self.margin, y, self.page_w - self.margin, y)
        p.set_y(y + 7)

    def code(self, code):
        self.check_page(12)
        self.pdf.set_font("Cou", "", 9)
        self.pdf.set_text_color(*self.c_code)
        lines = code.split("\n")
        # Fondo del bloque de código
        y0 = self.pdf.get_y()
        self.pdf.set_fill_color(*self.c_code_bg)
        lh = 4.5
        h = lh * len(lines) + 2
        if self.pdf.get_y() + h > 287:
            self.pdf.add_page()
            y0 = self.pdf.get_y()
        self.pdf.rect(self.margin, y0, self.content_w, h, "F")
        for i, line in enumerate(lines):
            self.pdf.set_xy(self.margin + 4, y0 + 2 + i * lh)
            self.pdf.cell(self.content_w - 8, lh, line)
        self.pdf.set_y(y0 + h + 1)

    def respuesta(self, lines, title="RESPUESTA / SOLUCIÓN"):
        """Bloque de respuesta con recuadro verde de fondo para diferenciarlo."""
        p = self.pdf
        lh = 5.5
        pad = 2

        def n_lines(text, size, style, w):
            p.set_font("Ar", style, size)
            return len(p.multi_cell(w, lh, text, dry_run=True, output="LINES"))

        t_w = self.content_w - 10  # indent 1.0
        r_w = self.content_w - 12  # indent 1.2
        h_title = n_lines(title, 11, "B", t_w) * lh
        h_resp = sum(n_lines(l, 10.5, "", r_w) for l in lines) * lh
        total = pad * 2 + h_title + 1.5 + h_resp

        gap = 6 * lh  # 6 espacios en blanco antes del bloque de respuesta
        # Si no cabe (incluyendo el hueco), saltar a página nueva para no partir el recuadro
        if p.get_y() + gap + total > 282:
            p.add_page()
        p.ln(gap)

        y0 = p.get_y()
        p.set_draw_color(*self.c_green)
        p.set_fill_color(*self.c_green_bg)
        p.rect(self.margin, y0, self.content_w, total, style="DF")
        p.set_y(y0 + pad)
        self.para(title, size=11, style="B", color=self.c_green, indent=1.0)
        for l in lines:
            self.para(l, size=10.5, color=self.c_green, indent=1.2)
        p.set_y(y0 + total)

    def title(self, text, size=20):
        self.pdf.add_page()
        self.pdf.set_text_color(*self.c_title)
        self.pdf.set_font("Ar", "B", size)
        self.pdf.set_xy(self.margin, self.pdf.get_y())
        self.pdf.multi_cell(self.content_w, 8, text, align="C")
        self.pdf.set_text_color(*self.c_text)

    def build(self, out_path):
        ex = self.exam
        p = self.pdf

        # --- Portada ---
        p.add_page()
        p.set_font("Ar", "B", 20)
        p.set_text_color(*self.c_title)
        p.set_x(self.margin)
        p.multi_cell(self.content_w, 9, ex["titulo"], align="C")
        p.set_font("Ar", "", 13)
        p.set_text_color(*self.c_sub)
        p.set_x(self.margin)
        p.ln(2)
        p.multi_cell(self.content_w, 7, ex["subtitulo"], align="C")
        p.set_font("Ar", "I", 10)
        p.set_text_color(*self.c_muted)
        p.set_x(self.margin)
        p.ln(3)
        p.multi_cell(self.content_w, 5.5, "\n".join(ex["temas"]), align="C")
        p.set_text_color(*self.c_text)
        p.ln(6)

        # Tabla informativa
        rows = [
            ("Asignatura", ASIGNATURA),
            ("Curso", CURSO),
            ("Duración", DURACION),
            ("Puntuación total", f'{ex["puntuacion"]} puntos'),
        ] + ex.get("info", [])
        p.set_font("Ar", "", 10)
        for k, v in rows:
            y = p.get_y()
            p.set_fill_color(*self.c_table)
            p.rect(self.margin, y, 60, 8, "F")
            p.set_xy(self.margin + 2, y + 1)
            p.set_font("Ar", "B", 10)
            p.cell(56, 6, k)
            p.set_xy(self.margin + 62, y + 1)
            p.set_font("Ar", "", 10)
            p.cell(self.content_w - 62, 6, v)
            p.ln(9)

        # --- Secciones ---
        seq = ex["secciones"]
        for idx, sec in enumerate(seq):
            p.add_page()
            p.set_font("Ar", "B", 15)
            p.set_text_color(*self.c_title)
            p.set_x(self.margin)
            p.multi_cell(self.content_w, 7, sec["titulo"])
            p.set_font("Ar", "I", 10)
            p.set_text_color(*self.c_muted)
            p.set_x(self.margin)
            p.multi_cell(self.content_w, 5, sec["puntos"])
            p.set_text_color(*self.c_text)
            p.ln(2)

            for num, q in enumerate(sec["preguntas"], start=1):
                enun = q.get("enunciado", "")
                if "puntos" in q:
                    enun = f"{enun}   ({q['puntos']})"
                # Número del ejercicio en rosa pastel + enunciado en negrita
                self.check_page(10)
                x0 = self.margin + 5  # indent 0.5
                self.pdf.set_xy(x0, self.pdf.get_y())
                self.pdf.set_font("Ar", "B", 11)
                self.pdf.set_text_color(*self.c_pink)
                self.pdf.cell(12, 5.5, f"{num}.")
                self.pdf.set_text_color(*self.c_text)
                self.pdf.set_xy(x0 + 12, self.pdf.get_y())
                self.pdf.multi_cell(self.content_w - 17, 5.5, enun, align="L")
                self.pdf.ln(1.5)

                if q.get("code"):
                    self.code(q["code"])

                if q.get("opciones"):
                    for o in q["opciones"]:
                        self.para(o, size=11, indent=1.2)

                if q.get("respuesta"):
                    resp = q["respuesta"] if isinstance(q["respuesta"], list) else [q["respuesta"]]
                    self.respuesta(resp)
                self.separador()

        # --- Plantilla de respuestas ---
        p.add_page()
        p.set_font("Ar", "B", 15)
        p.set_text_color(*self.c_title)
        p.set_x(self.margin)
        p.multi_cell(self.content_w, 7, "Plantilla de respuestas")
        p.set_font("Ar", "I", 10)
        p.set_text_color(*self.c_muted)
        p.set_x(self.margin)
        p.multi_cell(self.content_w, 5, "(Páginas en blanco para desarrollar las soluciones)")
        p.set_text_color(*self.c_text)
        p.ln(3)
        for _ in range(14):
            p.set_draw_color(*self.c_line)
            p.set_xy(self.margin, p.get_y())
            p.line(self.margin, p.get_y(), self.page_w - self.margin, p.get_y())
            p.ln(14)

        p.output(out_path)
        return out_path


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    out_dir = os.path.dirname(os.path.abspath(__file__))
    args = sys.argv[1:] if len(sys.argv) > 1 else ["docx", "pdf"]
    want_docx = "docx" in args
    want_pdf = "pdf" in args
    dark = "dark" in args or "oscuro" in args

    print(f"Directorios con fuentes disponibles: {bool(find_font('arial.ttf'))}")
    print("Modo PDF:", "OSCURO" if dark else "claro")

    for i, ex in enumerate(EXAMENES, start=1):
        nombre = f"Examen_{i}_Java_{ex['subtitulo'].split('·')[-1].strip().replace(' ','_')}"
        nombre = nombre.replace("/", "-").replace("\\", "-")[:80]
        base = os.path.join(out_dir, nombre)
        print(f"\n=== Examen {i}: {ex['subtitulo']} ===")

        if want_docx:
            path = base + ".docx"
            try:
                render_docx(ex, path, dark=dark)
                print("   DOCX OK:", os.path.basename(path))
            except Exception as e:
                print("   DOCX ERROR:", e)

        if want_pdf:
            path = base + ".pdf"
            try:
                PDFExam(ex, dark=dark).build(path)
                print("   PDF  OK:", os.path.basename(path))
            except Exception as e:
                import traceback
                print("   PDF  ERROR:", e)
                traceback.print_exc()

    print("\nFinalizado.")


if __name__ == "__main__":
    main()
