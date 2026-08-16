# -*- coding: utf-8 -*-
"""
Generador de 5 exámenes adicionales de Programación (Java) para Grado Superior.
Exámenes 6-10, con temas y ejercicios NUEVOS (sin repetir los de los exámenes 1-5).

Reutiliza el renderizado PDF de generar_examenes.py y usa una copia corregida
del render DOCX (etiqueta "RESPUESTA / SOLUCIÓN" única y código sin escapes).

Uso:
  python generar_examenes_2.py            # genera DOCX y PDF de los 5 exámenes
  python generar_examenes_2.py docx       # solo DOCX
  python generar_examenes_2.py pdf        # solo PDF
"""

import os
import sys

from generar_examenes import PDFExam, find_font, CURSO, ASIGNATURA, DURACION

# ---------------------------------------------------------------------------
# RENDERIZADO A DOCX (copiado de generar_examenes.py con dos correcciones:
# 1) la etiqueta "RESPUESTA / SOLUCIÓN" solo se añade una vez;
# 2) el código se inserta tal cual, python-docx escapa el XML por sí mismo).
# ---------------------------------------------------------------------------

def render_docx(exam, out_path, dark=False):
    from docx import Document
    from docx.shared import Pt, Cm, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Paleta: clara u oscura (mismos colores que el PDF)
    if dark:
        PAL = dict(bg="181A21", text="E2E6F0", title="81B2FF", sub="AAAFBE",
                   muted="969BAA", code="CED4E2", code_bg="262A36", table="282C3A",
                   green="7AD090", green_bg="1A2A1E", line="5F6473", pink="FFB6C1")
        # Fondo de página oscuro
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        bg = OxmlElement("w:background")
        bg.set(qn("w:color"), PAL["bg"])
        doc.settings.element.append(bg)
    else:
        PAL = dict(bg="FFFFFF", text="000000", title="1A478A", sub="555555",
                   muted="8C8C8C", code="323232", code_bg="F5F7FA", table="EEF3FA",
                   green="1B5E20", green_bg="E8F5E9", line="B4B4B4", pink="FFB6C1")

    def C(h):
        return RGBColor.from_string(h)

    for section in doc.sections:
        # Misma página A4 que el PDF
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = C(PAL["text"])

    AZUL = C(PAL["title"])
    VERDE = C(PAL["green"])
    ROSA = C(PAL["pink"])
    MUTED = C(PAL["muted"])
    SUB = C(PAL["sub"])

    def h1(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = AZUL
        return p

    def h2(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = False
        r.font.size = Pt(12)
        r.font.color.rgb = SUB
        return p

    def section_title(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = AZUL
        p.paragraph_format.space_before = Pt(14)
        return p

    def table_borders(t, color):
        """Color de todos los bordes de una tabla."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tblPr = t._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            borders.append(el)
        tblPr.append(borders)

    def info_table(rows):
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = "Table Grid"
        t.autofit = True
        for i, (k, v) in enumerate(rows):
            c0 = t.rows[i].cells[0]
            c1 = t.rows[i].cells[1]
            c0.text = k
            c1.text = v
            for p in c0.paragraphs:
                for r in p.runs:
                    r.bold = True
            shade_cell(c0, PAL["table"])  # azul como en el PDF
        table_borders(t, PAL["line"])
        return t

    def add_code(code):
        for line in code.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line)
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = C(PAL["code"])
            shade(p, PAL["code_bg"])  # fondo como el recuadro del PDF
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)
        sp.add_run("").font.size = Pt(2)

    def add_par(text, indent=0.0, italic=False, color=None, bold=False):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        r = p.add_run(text)
        r.italic = italic if italic else None
        r.bold = bold
        if color:
            r.font.color.rgb = color
        return p

    def shade(p, fill="E8F5E9"):
        """Sombreado de fondo para un párrafo (igual que el PDF)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        pPr.append(shd)

    def shade_cell(cell, fill):
        """Fondo de una celda de tabla."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    def cell_border(cell, color, sz=8):
        """Borde completo de celda (sz en octavos de punto)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            borders.append(el)
        tcPr.append(borders)

    def cell_margins(cell, top=80, left=120, bottom=80, right=120):
        """Márgenes internos de una celda (en veinteavos de punto)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tcPr = cell._tc.get_or_add_tcPr()
        mar = OxmlElement("w:tcMar")
        for edge, w in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:w"), str(w))
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        tcPr.append(mar)

    def p_border_bottom(p, color, sz=6):
        """Línea horizontal inferior en un párrafo (separadores del PDF)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(sz))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        pbdr.append(bottom)
        pPr.append(pbdr)

    # --- Portada ---
    h1(exam["titulo"])
    h2(exam["subtitulo"])
    add_par(" ".join(exam["temas"]), indent=0.0, italic=True, color=MUTED)
    doc.add_paragraph()

    info_rows = [
        ("Asignatura", ASIGNATURA),
        ("Curso", CURSO),
        ("Duración", DURACION),
        ("Puntuación total", f'{exam["puntuacion"]} puntos'),
    ]
    info_rows += exam["info"]
    info_table(info_rows)

    # --- Secciones ---
    for i, sec in enumerate(exam["secciones"]):
        if i > 0:
            doc.add_page_break()
        section_title(sec["titulo"])
        add_par(sec["puntos"], indent=0.0, italic=True, color=MUTED)

        for num, q in enumerate(sec["preguntas"], start=1):
            # Enunciado: número en rosa pastel + texto en negrita (igual que el PDF)
            enun = q.get("enunciado", "")
            if "puntos" in q:
                enun = f"{enun}  ({q['puntos']})"
            p_q = doc.add_paragraph()
            p_q.paragraph_format.left_indent = Cm(0.5)
            rn = p_q.add_run(f"{num}. ")
            rn.bold = True
            rn.font.size = Pt(11)
            rn.font.color.rgb = ROSA
            re_ = p_q.add_run(enun)
            re_.bold = True
            re_.font.size = Pt(11)

            if q.get("code"):
                add_code(q["code"])

            if q.get("opciones"):
                for o in q["opciones"]:
                    add_par(o, indent=1.2)

            if q.get("respuesta"):
                for _ in range(6):
                    doc.add_paragraph()
                resp = q["respuesta"] if isinstance(q["respuesta"], list) else [q["respuesta"]]
                # Recuadro verde con borde (igual que el bloque del PDF)
                t = doc.add_table(rows=1, cols=1)
                t.autofit = True
                cell = t.rows[0].cells[0]
                shade_cell(cell, PAL["green_bg"])
                cell_border(cell, PAL["green"], sz=8)
                cell_margins(cell)
                p0 = cell.paragraphs[0]
                r = p0.add_run("RESPUESTA / SOLUCIÓN")
                r.bold = True
                r.font.color.rgb = VERDE
                for l in resp:
                    p2 = cell.add_paragraph()
                    r2 = p2.add_run(l)
                    r2.font.color.rgb = VERDE
            # Línea separadora rosa pastel (igual que el PDF)
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(14)
            sep.paragraph_format.space_after = Pt(4)
            p_border_bottom(sep, "FFB6C1", sz=6)

    doc.add_page_break()
    section_title("Plantilla de respuestas")
    add_par("(Páginas en blanco para desarrollar las soluciones)", italic=True, color=MUTED)
    for _ in range(14):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p_border_bottom(p, PAL["line"], sz=6)

    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# CONTENIDO DE LOS 5 EXÁMENES NUEVOS (6-10)
# ---------------------------------------------------------------------------

EXAMENES = []

# =====================================================================
# EXAMEN 6 — Recursividad y algoritmos recursivos
# =====================================================================
EXAMENES.append({
    "titulo": "EXAMEN DE PROGRAMACIÓN – 1º GRADO SUPERIOR",
    "subtitulo": "Examen 6 · Recursividad y algoritmos recursivos",
    "temas": [
        "Recursión: concepto, caso base y caso recursivo",
        "La pila de llamadas y el desbordamiento de pila (StackOverflowError)",
        "Recursión directa, indirecta y recursión de cola",
        "Algoritmos recursivos clásicos: factorial, Fibonacci, MCD de Euclides",
        "Búsqueda binaria recursiva y Torres de Hanoi",
    ],
    "puntuacion": 16,
    "info": [("Convocatoria", "Evaluación 2"), ("Valoración", "Teoría y práctica de algoritmos recursivos")],
    "secciones": [
        {
            "titulo": "Parte 1: Opción múltiple y Verdadero/Falso",
            "puntos": "4 puntos — 0,5 cada pregunta",
            "preguntas": [
                {
                    "tipo": "test",
                    "enunciado": "Dado el siguiente método recursivo, ¿qué devuelve fact(4)?",
                    "code": """static int fact(int n) {
    if (n <= 1) return 1;
    return n * fact(n - 1);
}""",
                    "opciones": ["a) 12", "b) 24", "c) 4", "d) 120"],
                    "respuesta": "b) 24 — fact(4) = 4·fact(3) = 4·3·fact(2) = 4·3·2·fact(1) = 4·3·2·1 = 24. El caso base (n <= 1) detiene la recursión.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Con el método Fibonacci clásico (fib(0)=0, fib(1)=1, fib(n)=fib(n-1)+fib(n-2)), ¿cuánto vale fib(6)?",
                    "opciones": ["a) 5", "b) 6", "c) 8", "d) 13"],
                    "respuesta": "c) 8 — La sucesión es 0, 1, 1, 2, 3, 5, 8; por tanto fib(6) = 8.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "Toda función recursiva debe tener al menos un caso base; sin él se produce un desbordamiento de pila (StackOverflowError).",
                    "respuesta": "Verdadero — Sin caso base la función se llama a sí misma indefinidamente y agota la memoria de la pila de llamadas, lanzando StackOverflowError.",
                },
                {
                    "tipo": "test",
                    "enunciado": "En las Torres de Hanoi, ¿cuántos movimientos son necesarios como mínimo para mover una torre de 3 discos?",
                    "opciones": ["a) 3", "b) 5", "c) 7", "d) 9"],
                    "respuesta": "c) 7 — El número mínimo de movimientos es 2ⁿ − 1; para n = 3 discos: 2³ − 1 = 7.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Usando el algoritmo de Euclides recursivo, ¿qué devuelve mcd(48, 18)?",
                    "code": """static int mcd(int a, int b) {
    if (b == 0) return a;
    return mcd(b, a % b);
}""",
                    "opciones": ["a) 3", "b) 6", "c) 9", "d) 12"],
                    "respuesta": "b) 6 — mcd(48,18) → mcd(18,12) → mcd(12,6) → mcd(6,0) = 6.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Con la búsqueda binaria recursiva sobre el array ordenado {2, 5, 8, 12, 19, 25}, ¿qué índice devuelve buscar(19)?",
                    "opciones": ["a) 3", "b) 4", "c) 5", "d) -1"],
                    "respuesta": "b) 4 — El valor 19 está en la posición 4 del array (los índices empiezan en 0).",
                },
                {
                    "tipo": "vf",
                    "enunciado": "En la recursión de cola (tail recursion), la llamada recursiva es la última operación que ejecuta el método, lo que permite en algunos lenguajes optimizar la llamada reutilizando el marco de pila.",
                    "respuesta": "Verdadero — En la recursión de cola no queda trabajo pendiente después de la llamada recursiva; algunos compiladores/VMs la optimizan (aunque Java no la optimiza de forma general).",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué devuelve sumaDig(345) con el siguiente método?",
                    "code": """static int sumaDig(int n) {
    if (n < 10) return n;
    return n % 10 + sumaDig(n / 10);
}""",
                    "opciones": ["a) 12", "b) 15", "c) 8", "d) 345"],
                    "respuesta": "a) 12 — 5 + sumaDig(34) = 5 + 4 + sumaDig(3) = 5 + 4 + 3 = 12.",
                },
            ],
        },
        {
            "titulo": "Parte 2: Análisis y depuración",
            "puntos": "3 puntos — 1,5 cada ejercicio",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "El siguiente programa tiene 2 errores. Encuéntralos, explica qué ocurre al ejecutarlo y escribe la versión corregida.",
                    "puntos": "1,5 ptos",
                    "code": """public class PotenciaBug {
    static int potencia(int base, int exp) {
        return base * potencia(base, exp - 1);
    }
    public static void main(String[] args) {
        System.out.println(potencia(2, 3));
    }
}""",
                    "respuesta": [
                        "1. Falta el caso base: la función se llama a sí misma sin condición de parada, por lo que la recursión es infinita y el programa termina con StackOverflowError.",
                        "2. Tampoco se contempla el caso exp == 0 (debería devolver 1) ni los exponentes negativos, que nunca llegarían a 0.",
                        "Versión corregida:",
                        """static int potencia(int base, int exp) {
    if (exp == 0) return 1;                 // caso base
    return base * potencia(base, exp - 1);  // caso recursivo
}""",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Analiza el siguiente método y realiza la traza de la ejecución de fib(5). Indica el valor devuelto y el número total de llamadas a fib que se realizan.",
                    "puntos": "1,5 ptos",
                    "code": """static int fib(int n) {
    if (n < 2) return n;
    return fib(n - 1) + fib(n - 2);
}""",
                    "respuesta": [
                        "Valores: fib(0)=0, fib(1)=1, fib(2)=1, fib(3)=2, fib(4)=3, fib(5)=5.",
                        "Árbol de llamadas de fib(5): fib(5) llama a fib(4) y fib(3); fib(4) llama a fib(3) y fib(2); fib(3) llama a fib(2) y fib(1); fib(2) llama a fib(1) y fib(0).",
                        "Total de llamadas: 15 (1 de fib(5) + 2 del nivel 2 + 4 del nivel 3 + 8 del nivel 4).",
                        "Resultado: fib(5) = 5.",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 3: Ejercicios de programación",
            "puntos": "8 puntos — 2 cada ejercicio",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "Implementa un método recursivo potencia(int base, int exp) que calcule base^exp para exponentes mayores o iguales que 0. Escribe también un main que pruebe 2^10, 3^0 y 5^3.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        """public class Potencia {
    static long potencia(int base, int exp) {
        if (exp == 0) return 1;                // caso base
        return base * potencia(base, exp - 1); // caso recursivo
    }
    public static void main(String[] args) {
        System.out.println("2^10 = " + potencia(2, 10));  // 1024
        System.out.println("3^0  = " + potencia(3, 0));   // 1
        System.out.println("5^3  = " + potencia(5, 3));   // 125
    }
}""",
                        "Nota: para exponentes negativos haría falta trabajar con double (base^-n = 1 / base^n).",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Escribe una función recursiva que devuelva una cadena invertida y un main que la pruebe con la cadena «Hola mundo».",
                    "puntos": "2 ptos",
                    "respuesta": [
                        """public class InvertirCadena {
    static String invertir(String s) {
        if (s.length() <= 1) return s;               // caso base
        return invertir(s.substring(1)) + s.charAt(0);
    }
    public static void main(String[] args) {
        String texto = "Hola mundo";
        System.out.println(invertir(texto));         // "odnum aloH"
    }
}""",
                        "Explicación: invertir(\"Hola\") = invertir(\"ola\") + 'H' = (invertir(\"la\") + 'o') + 'H' = ... = \"aloH\".",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Implementa la solución recursiva de las Torres de Hanoi para n discos: el método debe imprimir cada movimiento («Mover disco X de A a C») y el main debe ejecutarlo con 3 discos mostrando además el número total de movimientos.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        """public class Hanoi {
    static int movimientos = 0;

    static void hanoi(int n, char origen, char auxiliar, char destino) {
        if (n == 1) {
            System.out.println("Mover disco 1 de " + origen + " a " + destino);
            movimientos++;
        } else {
            hanoi(n - 1, origen, destino, auxiliar);
            System.out.println("Mover disco " + n + " de " + origen + " a " + destino);
            movimientos++;
            hanoi(n - 1, auxiliar, origen, destino);
        }
    }
    public static void main(String[] args) {
        hanoi(3, 'A', 'B', 'C');
        System.out.println("Total de movimientos: " + movimientos);  // 7
    }
}""",
                        "El número mínimo de movimientos es 2ⁿ − 1; para 3 discos se imprimen 7 movimientos.",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Implementa la búsqueda binaria de forma recursiva sobre un array ordenado de enteros. Escribe un main que busque el 19 y el 7 en {2, 5, 8, 12, 19, 25} mostrando el índice encontrado o «No encontrado».",
                    "puntos": "2 ptos",
                    "respuesta": [
                        """public class BusquedaBinaria {
    static int buscar(int[] a, int ini, int fin, int x) {
        if (ini > fin) return -1;                 // caso base: no encontrado
        int medio = (ini + fin) / 2;
        if (a[medio] == x) return medio;          // caso base: encontrado
        if (x < a[medio]) return buscar(a, ini, medio - 1, x);
        return buscar(a, medio + 1, fin, x);
    }
    public static void main(String[] args) {
        int[] a = {2, 5, 8, 12, 19, 25};
        int pos = buscar(a, 0, a.length - 1, 19);
        System.out.println(pos >= 0 ? "Índice: " + pos : "No encontrado");   // 4
        pos = buscar(a, 0, a.length - 1, 7);
        System.out.println(pos >= 0 ? "Índice: " + pos : "No encontrado");   // No encontrado
    }
}""",
                        "Cada llamada divide el intervalo a la mitad, por lo que la complejidad es O(log n).",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 4: Pregunta teórica",
            "puntos": "1 punto",
            "preguntas": [
                {
                    "tipo": "teorica",
                    "enunciado": "Explica qué es la recursión, qué papel juegan el caso base y la pila de llamadas, y en qué situaciones conviene usar recursión en lugar de iteración. ¿Qué riesgo tiene un uso incorrecto?",
                    "respuesta": [
                        "La recursión es una técnica en la que un método se llama a sí mismo para resolver un problema dividiéndolo en versiones más pequeñas del mismo problema.",
                        "El caso base es la condición que detiene las llamadas (p. ej., n <= 1 en el factorial); sin él la recursión es infinita. Cada llamada pendiente ocupa un marco en la pila de llamadas (stack), que guarda sus parámetros y variables locales; al llegar al caso base, los marcos se desapilan devolviendo resultados.",
                        "Conviene usar recursión cuando el problema se define de forma natural recursiva (árboles, directorios, Hanoi, divide y vencerás) y cuando la versión iterativa sería mucho más compleja. La iteración es preferible cuando la profundidad puede ser grande, porque la recursión consume memoria de pila y un uso incorrecto (sin caso base o con demasiada profundidad) lanza StackOverflowError.",
                    ],
                },
            ],
        },
    ],
})

# =====================================================================
# EXAMEN 7 — Clases y objetos avanzados: enum, wrappers, StringBuilder,
#             equals y hashCode
# =====================================================================
EXAMENES.append({
    "titulo": "EXAMEN DE PROGRAMACIÓN – 1º GRADO SUPERIOR",
    "subtitulo": "Examen 7 · POO avanzada: enum, wrappers, StringBuilder, equals y hashCode",
    "temas": [
        "Tipos enumerados (enum): valores, ordinal(), name() y valueOf()",
        "Clases envoltorio (wrappers): autoboxing, unboxing y caché de Integer",
        "String vs StringBuilder vs StringBuffer: inmutabilidad y eficiencia",
        "El contrato equals() y hashCode()",
        "Composición y agregación entre clases. Objetos inmutables",
    ],
    "puntuacion": 16,
    "info": [("Convocatoria", "Evaluación 2"), ("Valoración", "POO avanzada y clases de utilidad")],
    "secciones": [
        {
            "titulo": "Parte 1: Opción múltiple y Verdadero/Falso",
            "puntos": "4 puntos — 0,5 cada pregunta",
            "preguntas": [
                {
                    "tipo": "test",
                    "enunciado": "Dado el enum: enum Dia { LUNES, MARTES, MIÉRCOLES, JUEVES }, ¿qué devuelve Dia.MARTES.ordinal()?",
                    "opciones": ["a) 0", "b) 1", "c) 2", "d) Error de compilación"],
                    "respuesta": "b) 1 — ordinal() devuelve la posición del valor dentro del enum empezando en 0 (LUNES=0, MARTES=1).",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué imprime este código?",
                    "code": """Integer a = 500;
Integer b = 500;
System.out.println(a == b);""",
                    "opciones": ["a) true", "b) false", "c) Error de compilación", "d) NullPointerException"],
                    "respuesta": "b) false — a == b compara referencias, y el valor 500 está fuera de la caché de Integer (-128..127), por lo que son objetos distintos. Para comparar valores debe usarse a.equals(b) o a.intValue() == b.intValue().",
                },
                {
                    "tipo": "vf",
                    "enunciado": "StringBuilder es una clase inmutable: una vez creado un objeto StringBuilder no se puede modificar su contenido.",
                    "respuesta": "Falso — StringBuilder es mutable: métodos como append(), insert() o reverse() modifican el mismo objeto. La clase inmutable es String.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Con estas declaraciones, ¿qué línea imprime true?",
                    "code": """String s1 = new String("hola");
String s2 = new String("hola");
System.out.println(s1 == s2);       // línea 1
System.out.println(s1.equals(s2));  // línea 2""",
                    "opciones": ["a) Solo la línea 1", "b) Solo la línea 2", "c) Ambas líneas", "d) Ninguna"],
                    "respuesta": "b) Solo la línea 2 — == compara referencias (s1 y s2 son objetos distintos → false), mientras que equals() compara el contenido («hola» == «hola» → true).",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué imprime este código?",
                    "code": """Integer a = 100;
Integer b = 100;
System.out.println(a == b);""",
                    "opciones": ["a) true", "b) false", "c) Error de compilación", "d) 100"],
                    "respuesta": "a) true — Los valores entre -128 y 127 están en la caché de Integer, por lo que a y b referencian el mismo objeto y == devuelve true.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "Si una clase sobrescribe equals(), debe sobrescribir también hashCode() para cumplir el contrato: dos objetos iguales según equals() deben tener el mismo hashCode().",
                    "respuesta": "Verdadero — Si no se cumple, estructuras como HashMap o HashSet pueden fallar (objetos iguales en cubetas distintas).",
                },
                {
                    "tipo": "test",
                    "enunciado": "Un objeto Coche crea en su constructor un objeto Motor y este no puede existir sin el coche. La relación entre Coche y Motor se llama…",
                    "opciones": ["a) Herencia", "b) Composición", "c) Polimorfismo", "d) Agregación"],
                    "respuesta": "b) Composición — Relación fuerte «tiene-un» (has-a) en la que el ciclo de vida del Motor está ligado al del Coche. En la agregación el objeto se crea fuera y se pasa al contenedor.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué imprime este código?",
                    "code": """StringBuilder sb = new StringBuilder("abc");
sb.append("def").insert(3, "-");
System.out.println(sb);""",
                    "opciones": ["a) abc-def", "b) abcdef-", "c) -abcdef", "d) abc-def-"],
                    "respuesta": "a) abc-def — append(\"def\") deja «abcdef» y insert(3, \"-\") coloca el guion en el índice 3: «abc-def».",
                },
            ],
        },
        {
            "titulo": "Parte 2: Análisis y depuración",
            "puntos": "3 puntos — 1,5 cada ejercicio",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "El siguiente programa tiene 2 errores. Encuéntralos, explica qué ocurre en cada uno y escribe la versión corregida.",
                    "puntos": "1,5 ptos",
                    "code": """public class WrapperBug {
    public static void main(String[] args) {
        Integer n = null;
        int x = n;                  // error 1
        Integer a = 500;
        Integer b = 500;
        System.out.println(a == b); // error 2
    }
}""",
                    "respuesta": [
                        "1. `int x = n;` con n = null lanza NullPointerException: el unboxing de un Integer null no está permitido. Solución: comprobar n != null antes o no asignar null.",
                        "2. `a == b` compara referencias y no valores: con 500 (fuera de la caché) imprime false. Solución: usar a.equals(b) o a.intValue() == b.intValue().",
                        """Versión corregida:
    Integer a = 500;
    Integer b = 500;
    System.out.println(a.equals(b));   // true
    Integer n = 10;
    int x = n;                          // unboxing válido: n no es null
    System.out.println(x);""",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Realiza la traza del siguiente código e indica qué se imprime en cada línea, explicando la diferencia entre String y StringBuilder.",
                    "puntos": "1,5 ptos",
                    "code": """StringBuilder sb = new StringBuilder("Java");
sb.reverse();
sb.append(" 8");
System.out.println(sb);

String s = "Java";
String s2 = s.toUpperCase();
System.out.println(s);
System.out.println(s2);""",
                    "respuesta": [
                        "sb.reverse() invierte el contenido del MISMO objeto: «avaJ»; sb.append(\" 8\") añade al final: «avaJ 8». Se imprime: avaJ 8.",
                        "String es inmutable: toUpperCase() NO modifica s, sino que devuelve un objeto nuevo. Se imprime «Java» (s intacto) y «JAVA» (s2).",
                        "Conclusión: StringBuilder modifica el objeto en el sitio (mutable); String crea un objeto nuevo en cada operación.",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 3: Ejercicios de programación",
            "puntos": "8 puntos — 2 cada ejercicio",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "Crea la clase CuentaBancaria con un enum interno TipoCuenta { AHORRO, CORRIENTE }, los atributos privados numero (String), saldo (double) y tipo, y los métodos getNumero(), getSaldo(), getTipo() e ingresar(double). Sobrescribe equals() y hashCode() para que dos cuentas sean iguales si tienen el mismo número. Escribe un main que cree dos cuentas con el mismo número y distintos tipos y compruebe que equals() devuelve true.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        """import java.util.Objects;

public class CuentaBancaria {
    enum TipoCuenta { AHORRO, CORRIENTE }

    private String numero;
    private double saldo;
    private TipoCuenta tipo;

    public CuentaBancaria(String numero, TipoCuenta tipo) {
        this.numero = numero;
        this.tipo = tipo;
        this.saldo = 0;
    }
    public String getNumero() { return numero; }
    public double getSaldo() { return saldo; }
    public TipoCuenta getTipo() { return tipo; }

    public void ingresar(double cantidad) { saldo += cantidad; }

    @Override
    public boolean equals(Object o) {
        if (this == o) return true;
        if (!(o instanceof CuentaBancaria)) return false;
        CuentaBancaria c = (CuentaBancaria) o;
        return numero.equals(c.numero);
    }
    @Override
    public int hashCode() {
        return Objects.hash(numero);
    }

    public static void main(String[] args) {
        CuentaBancaria c1 = new CuentaBancaria("ES1234", TipoCuenta.AHORRO);
        CuentaBancaria c2 = new CuentaBancaria("ES1234", TipoCuenta.CORRIENTE);
        c1.ingresar(100);
        System.out.println(c1.equals(c2));          // true: mismo número
        System.out.println(c1.getSaldo());          // 100.0
        System.out.println(TipoCuenta.AHORRO.ordinal());  // 0
    }
}""",
                        "equals() ignora saldo y tipo: la identidad de la cuenta es su número. hashCode() usa el mismo campo para mantener el contrato.",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Escribe un programa que lea una frase por teclado y la imprima con las palabras en orden inverso usando StringBuilder. Ejemplo: «hola mundo java» → «java mundo hola».",
                    "puntos": "2 ptos",
                    "respuesta": [
                        """import java.util.Scanner;

public class InvertirFrase {
    public static void main(String[] args) {
        Scanner sc = new Scanner(System.in);
        System.out.print("Frase: ");
        String frase = sc.nextLine();
        String[] palabras = frase.split(" ");
        StringBuilder sb = new StringBuilder();
        for (int i = palabras.length - 1; i >= 0; i--) {
            sb.append(palabras[i]);
            if (i > 0) sb.append(" ");
        }
        System.out.println("Invertida: " + sb);
    }
}""",
                        "Se recorre el array de palabras de atrás hacia adelante y se concatenan con StringBuilder, que es eficiente porque no crea un String nuevo en cada append.",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Implementa mediante composición las clases Motor (atributo caballos y método arrancar() que imprime «Motor de X CV en marcha») y Coche (marca y un Motor creado en su constructor). El método arrancar() del Coche debe delegar en el motor. Escribe un main que cree un Coche y lo arranque, y justifica por qué es composición y no herencia.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        """public class CocheMain {
    static class Motor {
        private int caballos;
        Motor(int caballos) { this.caballos = caballos; }
        void arrancar() {
            System.out.println("Motor de " + caballos + " CV en marcha");
        }
    }
    static class Coche {
        private String marca;
        private Motor motor;                  // composición: Coche TIENE-UN Motor
        Coche(String marca, int caballos) {
            this.marca = marca;
            this.motor = new Motor(caballos); // el Motor nace con el Coche
        }
        void arrancar() { motor.arrancar(); } // delega en el motor
    }
    public static void main(String[] args) {
        Coche c = new Coche("Seat", 110);
        c.arrancar();
    }
}""",
                        "Es composición porque Coche no «es un» Motor (no hereda), sino que lo contiene; además el Motor se crea dentro del constructor del Coche, de modo que su ciclo de vida está ligado al del Coche. La salida es: Motor de 110 CV en marcha.",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Escribe un programa que lea 5 números por teclado como String, los convierta con Integer.parseInt, los sume con autoboxing/unboxing (variable Integer total) y muestre el resultado. Debe capturar NumberFormatException si el texto no es un número y pedir el valor de nuevo.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        """import java.util.Scanner;

public class SumaNumeros {
    public static void main(String[] args) {
        Scanner sc = new Scanner(System.in);
        Integer total = 0;                     // autoboxing
        for (int i = 0; i < 5; i++) {
            System.out.print("Número " + (i + 1) + ": ");
            String texto = sc.nextLine();
            try {
                int n = Integer.parseInt(texto); // String -> int
                total = total + n;               // unboxing + autoboxing
            } catch (NumberFormatException e) {
                System.out.println("'" + texto + "' no es un número válido");
                i--;                             // repetir la lectura
            }
        }
        System.out.println("Suma: " + total);
    }
}""",
                        "Integer.parseInt convierte el texto; si falla lanza NumberFormatException (subclase de IllegalArgumentException, no comprobada), que se captura para repetir la lectura.",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 4: Pregunta teórica",
            "puntos": "1 punto",
            "preguntas": [
                {
                    "tipo": "teorica",
                    "enunciado": "Explica las diferencias entre String, StringBuilder y StringBuffer (inmutabilidad, eficiencia y sincronización) y en qué situaciones usarías cada uno.",
                    "respuesta": [
                        "String es inmutable: cualquier operación (concatenación, toUpperCase, etc.) crea un objeto nuevo. Es eficiente para valores que no cambian y para usarlo como clave en mapas.",
                        "StringBuilder es mutable y no sincronizado: modifica el mismo objeto, por lo que es la opción recomendada para concatenar muchas cadenas en un bucle (evita crear decenas de objetos intermedios).",
                        "StringBuffer es igual que StringBuilder pero con métodos sincronizados (thread-safe), lo que lo hace más lento; solo tiene sentido en programas multihilo donde varios hilos comparten el objeto.",
                        "Regla práctica: String para texto fijo, StringBuilder para construir cadenas dinámicamente en un solo hilo, StringBuffer solo si hay concurrencia.",
                    ],
                },
            ],
        },
    ],
})

# =====================================================================
# EXAMEN 8 — Colecciones y genéricos
# =====================================================================
EXAMENES.append({
    "titulo": "EXAMEN DE PROGRAMACIÓN – 1º GRADO SUPERIOR",
    "subtitulo": "Examen 8 · Colecciones y genéricos",
    "temas": [
        "La interfaz List: ArrayList y LinkedList. Iterator",
        "Conjuntos: HashSet, LinkedHashSet y TreeSet. Orden e igualdad",
        "Mapas: HashMap, LinkedHashMap y TreeMap",
        "Comparable vs Comparator. Ordenación de colecciones",
        "Genéricos: clases y métodos genéricos. Autoboxing en colecciones",
    ],
    "puntuacion": 16,
    "info": [("Convocatoria", "Evaluación 2"), ("Valoración", "Colecciones y genéricos")],
    "secciones": [
        {
            "titulo": "Parte 1: Opción múltiple y Verdadero/Falso",
            "puntos": "4 puntos — 0,5 cada pregunta",
            "preguntas": [
                {
                    "tipo": "test",
                    "enunciado": "¿Qué implementación de List es más eficiente para insertar y eliminar elementos en posiciones intermedias?",
                    "opciones": ["a) ArrayList", "b) LinkedList", "c) TreeSet", "d) HashMap"],
                    "respuesta": "b) LinkedList — La inserción/eliminación en medio solo requiere reenlazar nodos; en ArrayList hay que desplazar todos los elementos posteriores.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué imprime el siguiente código?",
                    "code": """TreeSet<Integer> ts = new TreeSet<>();
ts.add(5); ts.add(1); ts.add(5); ts.add(3);
System.out.println(ts.size() + " " + ts);""",
                    "opciones": ["a) 4 [1, 3, 5, 5]", "b) 3 [1, 3, 5]", "c) 3 [5, 3, 1]", "d) 4 [5, 3, 1, 5]"],
                    "respuesta": "b) 3 [1, 3, 5] — TreeSet no admite duplicados (el segundo 5 se descarta) y mantiene el orden natural ascendente.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Dado el código, ¿qué se imprime?",
                    "code": """Map<String, Integer> m = new HashMap<>();
m.put("a", 1);
System.out.println(m.put("a", 2));""",
                    "opciones": ["a) null", "b) 1", "c) 2", "d) a"],
                    "respuesta": "b) 1 — put() devuelve el valor anterior asociado a la clave (o null si no existía) y después lo sustituye por 2.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "HashSet permite almacenar elementos duplicados.",
                    "respuesta": "Falso — HashSet rechaza duplicados: si se añade un elemento igual (según equals/hashCode) a uno existente, add() devuelve false y no lo inserta.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Al recorrer un ArrayList con for-each y llamar a remove() dentro del bucle se lanza una excepción. ¿Cuál?",
                    "opciones": ["a) NullPointerException", "b) ConcurrentModificationException", "c) IndexOutOfBoundsException", "d) No lanza ninguna excepción"],
                    "respuesta": "b) ConcurrentModificationException — El iterador del for-each detecta que la lista se modificó y falla; hay que usar Iterator y su método remove().",
                },
                {
                    "tipo": "test",
                    "enunciado": "Para que una clase tenga un orden natural y pueda usarse en TreeSet o Collections.sort() debe implementar una interfaz. ¿Cuál?",
                    "opciones": ["a) Comparator con compare()", "b) Comparable con compareTo()", "c) Iterator con hasNext()", "d) Serializable"],
                    "respuesta": "b) Comparable con compareTo() — El orden natural se define implementando Comparable; Comparator permite definir órdenes alternativos sin tocar la clase.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Dada la clase genérica class Caja<T> { ... }, ¿qué tipo concreto tiene T en la declaración Caja<String> c = new Caja<>();?",
                    "opciones": ["a) String", "b) Object", "c) T", "d) Error de compilación"],
                    "respuesta": "a) String — El parámetro de tipo se instancia con String; el operador diamante <> infiere el tipo en el constructor.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "TreeMap mantiene sus claves ordenadas según el orden natural de la clave o el Comparator proporcionado.",
                    "respuesta": "Verdadero — TreeMap ordena las claves (orden natural o Comparator); HashMap no garantiza ningún orden y LinkedHashMap mantiene el orden de inserción.",
                },
            ],
        },
        {
            "titulo": "Parte 2: Análisis y depuración",
            "puntos": "3 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "El siguiente programa tiene 2 errores. Encuéntralos, explica por qué fallan y propón la corrección.",
                    "puntos": "1,5 ptos",
                    "code": """import java.util.*;
public class BorrarPares {
    public static void main(String[] args) {
        List<Integer> numeros = new ArrayList<>(List.of(1, 2, 3, 4, 5, 6));
        for (Integer n : numeros) {
            if (n % 2 == 0) numeros.remove(n);
        }
        Integer x = numeros.get(0);
        Integer y = numeros.get(1);
        System.out.println(x == y);
    }
}""",
                    "respuesta": [
                        "1. Modificar la lista con remove() durante un for-each lanza ConcurrentModificationException. Corrección: recorrer con Iterator y usar it.remove():",
                        "   Iterator<Integer> it = numeros.iterator();",
                        "   while (it.hasNext()) { if (it.next() % 2 == 0) it.remove(); }",
                        "2. x == y compara referencias de Integer, no valores: si ambos son iguales pero están fuera de la caché (-128..127) daría false. Corrección: x.equals(y) o x.intValue() == y.intValue().",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Analiza el siguiente código y determina qué orden imprime el programa (traza manual).",
                    "puntos": "1,5 ptos",
                    "code": """List<Alumno> lista = new ArrayList<>();
lista.add(new Alumno("Ana", 7.5));
lista.add(new Alumno("Luis", 5.0));
lista.add(new Alumno("Eva", 9.2));
lista.add(new Alumno("Juan", 6.3));
lista.sort((a1, a2) -> Double.compare(a2.nota, a1.nota));
for (Alumno a : lista) System.out.println(a.nombre);""",
                    "respuesta": [
                        "El comparador usa Double.compare(a2.nota, a1.nota): al invertir los operandos, el orden queda DESCENDENTE por nota.",
                        "Salida: Eva, Ana, Juan, Luis.",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 3: Ejercicios de programación",
            "puntos": "8 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "Escribe un programa que cuente cuántas veces aparece cada palabra en una frase y muestre el resultado con un HashMap<String, Integer>.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.util.*;",
                        "public class Frecuencias {",
                        "    public static void main(String[] args) {",
                        "        String frase = \"el sol brilla y el sol calienta\";",
                        "        String[] palabras = frase.split(\" \");",
                        "        Map<String, Integer> contador = new HashMap<>();",
                        "        for (String p : palabras) {",
                        "            contador.put(p, contador.getOrDefault(p, 0) + 1);",
                        "        }",
                        "        for (String p : contador.keySet()) {",
                        "            System.out.println(p + \" -> \" + contador.get(p));",
                        "        }",
                        "    }",
                        "}",
                        "Salida: el -> 2, sol -> 2, brilla -> 1, y -> 1, calienta -> 1 (el orden de HashMap puede variar).",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Implementa una clase genérica Par<K, V> que almacene una clave y un valor (con getClave(), getValor() y setValor()). Úsala en un main con Par<String, Integer> y Par<String, String>.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "public class Par<K, V> {",
                        "    private K clave;",
                        "    private V valor;",
                        "    public Par(K clave, V valor) { this.clave = clave; this.valor = valor; }",
                        "    public K getClave() { return clave; }",
                        "    public V getValor() { return valor; }",
                        "    public void setValor(V valor) { this.valor = valor; }",
                        "    public static void main(String[] args) {",
                        "        Par<String, Integer> p = new Par<>(\"edad\", 25);",
                        "        System.out.println(p.getClave() + \" = \" + p.getValor());",
                        "        p.setValor(26);",
                        "        System.out.println(p.getClave() + \" = \" + p.getValor());",
                        "        Par<String, String> q = new Par<>(\"ciudad\", \"Madrid\");",
                        "        System.out.println(q.getClave() + \" = \" + q.getValor());",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Crea una clase Alumno (nombre, nota) y un programa que guarde varios alumnos en un ArrayList, los ordene por nota descendente con Comparator y muestre también la nota media.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.util.*;",
                        "public class OrdenarAlumnos {",
                        "    static class Alumno {",
                        "        String nombre; double nota;",
                        "        Alumno(String nombre, double nota) { this.nombre = nombre; this.nota = nota; }",
                        "        public String toString() { return nombre + \" (\" + nota + \")\"; }",
                        "    }",
                        "    public static void main(String[] args) {",
                        "        List<Alumno> lista = new ArrayList<>();",
                        "        lista.add(new Alumno(\"Ana\", 7.5));",
                        "        lista.add(new Alumno(\"Luis\", 5.0));",
                        "        lista.add(new Alumno(\"Eva\", 9.2));",
                        "        lista.add(new Alumno(\"Juan\", 6.3));",
                        "        lista.sort((a1, a2) -> Double.compare(a2.nota, a1.nota));",
                        "        double suma = 0;",
                        "        for (Alumno a : lista) suma += a.nota;",
                        "        System.out.println(\"Media: \" + (suma / lista.size()));",
                        "        for (Alumno a : lista) System.out.println(a);",
                        "    }",
                        "}",
                        "Salida: Media: 7.0; Eva (9.2), Ana (7.5), Juan (6.3), Luis (5.0).",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Usa un TreeSet con un Comparator que ordene las palabras primero por longitud y después alfabéticamente. Inserta {\"sol\", \"mar\", \"cielo\", \"sol\", \"luna\"} y muestra el tamaño y el contenido final.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.util.*;",
                        "public class OrdenarLongitud {",
                        "    public static void main(String[] args) {",
                        "        Comparator<String> porLongitud = new Comparator<>() {",
                        "            public int compare(String a, String b) {",
                        "                if (a.length() != b.length()) return a.length() - b.length();",
                        "                return a.compareTo(b);",
                        "            }",
                        "        };",
                        "        TreeSet<String> set = new TreeSet<>(porLongitud);",
                        "        set.add(\"sol\"); set.add(\"mar\"); set.add(\"cielo\");",
                        "        set.add(\"sol\"); set.add(\"luna\");",
                        "        System.out.println(set.size());",
                        "        System.out.println(set);",
                        "    }",
                        "}",
                        "Salida: tamaño 4 (\"sol\" duplicado se descarta) y contenido [mar, sol, luna, cielo].",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 4: Pregunta teórica",
            "puntos": "1 punto",
            "preguntas": [
                {
                    "tipo": "teorica",
                    "enunciado": "Explica las diferencias entre List, Set y Map. ¿Qué implementación elegirías para: (a) acceso frecuente por índice, (b) evitar duplicados conservando el orden de inserción, (c) buscar valores por clave rápidamente? Razona la respuesta.",
                    "respuesta": [
                        "List: secuencia ordenada que admite duplicados y acceso por índice. Set: colección sin duplicados (igualdad por equals/hashCode). Map: pares clave-valor sin claves repetidas.",
                        "(a) ArrayList: acceso por índice O(1). (b) LinkedHashSet: sin duplicados y mantiene orden de inserción (HashSet no garantiza orden; TreeSet ordena, no conserva inserción). (c) HashMap: búsqueda por clave O(1) media gracias al hashCode.",
                    ],
                },
            ],
        },
    ],
})

# =====================================================================
# EXAMEN 9 — Ficheros binarios, serialización y NIO
# =====================================================================
EXAMENES.append({
    "titulo": "EXAMEN DE PROGRAMACIÓN – 1º GRADO SUPERIOR",
    "subtitulo": "Examen 9 · Ficheros binarios, serialización y NIO",
    "temas": [
        "Flujos de bytes: FileInputStream y FileOutputStream",
        "Flujos de datos: DataInputStream y DataOutputStream",
        "Serialización: ObjectOutputStream, ObjectInputStream, Serializable y transient",
        "Lectura/escritura de texto con buffer: BufferedReader y BufferedWriter",
        "java.nio.file: Path, Files (readAllLines, write, walk)",
    ],
    "puntuacion": 16,
    "info": [("Convocatoria", "Evaluación 3"), ("Valoración", "Ficheros binarios y serialización")],
    "secciones": [
        {
            "titulo": "Parte 1: Opción múltiple y Verdadero/Falso",
            "puntos": "4 puntos — 0,5 cada pregunta",
            "preguntas": [
                {
                    "tipo": "test",
                    "enunciado": "¿Qué interfaz debe implementar una clase para poder guardar sus objetos con ObjectOutputStream?",
                    "opciones": ["a) Cloneable", "b) Comparable", "c) Serializable", "d) Iterator"],
                    "respuesta": "c) Serializable — Es una interfaz de marcado (sin métodos) que indica que el objeto puede serializarse.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "Un campo declarado como transient no se guarda al serializar el objeto.",
                    "respuesta": "Verdadero — transient excluye el campo de la serialización; al deserializar queda con el valor por defecto (0, null, false...).",
                },
                {
                    "tipo": "test",
                    "enunciado": "Si guardamos un int con out.writeInt(7), ¿qué método hay que usar para leerlo en el mismo orden?",
                    "opciones": ["a) readDouble()", "b) readInt()", "c) readLine()", "d) readObject()"],
                    "respuesta": "b) readInt() — Los flujos de datos (DataInputStream/DataOutputStream) exigen leer los tipos en el mismo orden en que se escribieron.",
                },
                {
                    "tipo": "test",
                    "enunciado": "Con try-with-resources (try (BufferedReader br = ...) { ... }), al terminar el bloque, ¿qué ocurre con br?",
                    "opciones": ["a) hay que cerrar br manualmente", "b) br se cierra automáticamente", "c) br se borra del disco", "d) se lanza IOException siempre"],
                    "respuesta": "b) br se cierra automáticamente — try-with-resources cierra todos los recursos declarados entre paréntesis al finalizar el bloque, incluso si hay una excepción.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "Files.write(Path, List<String>) permite escribir un fichero de texto a partir de una lista de líneas.",
                    "respuesta": "Verdadero — java.nio.file.Files.write() crea (o sobrescribe) el fichero y escribe cada String de la lista como una línea.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué excepción lanza ObjectOutputStream al intentar serializar un objeto que NO implementa Serializable?",
                    "opciones": ["a) IOException", "b) NotSerializableException", "c) ClassNotFoundException", "d) FileNotFoundException"],
                    "respuesta": "b) NotSerializableException — Es una subclase de IOException que se lanza cuando un objeto no es serializable.",
                },
                {
                    "tipo": "test",
                    "enunciado": "El método read() de FileInputStream devuelve un valor. ¿Cuál?",
                    "opciones": ["a) un char", "b) un int con el byte leído (o -1 si es fin de fichero)", "c) un boolean", "d) un String con la línea"],
                    "respuesta": "b) un int con el byte leído (o -1 si es fin de fichero) — Por eso el bucle típico de copia es while ((b = in.read()) != -1).",
                },
                {
                    "tipo": "vf",
                    "enunciado": "Un fichero binario puede leerse con Scanner igual que un fichero de texto.",
                    "respuesta": "Falso — Scanner está pensado para texto; para datos binarios hay que usar flujos de bytes o de datos (FileInputStream, DataInputStream, ObjectInputStream...).",
                },
            ],
        },
        {
            "titulo": "Parte 2: Análisis y depuración",
            "puntos": "3 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "El siguiente programa tiene 2 errores. Encuéntralos, explica por qué fallan y propón la corrección.",
                    "puntos": "1,5 ptos",
                    "code": """public class Guardar {
    public static void main(String[] args) throws IOException {
        Persona p = new Persona("Ana", 22);
        ObjectOutputStream out =
            new ObjectOutputStream(new FileOutputStream("p.dat"));
        out.writeObject(p);
    }
}""",
                    "respuesta": [
                        "1. La clase Persona no implementa Serializable: al ejecutar out.writeObject(p) se lanza NotSerializableException. Corrección: public class Persona implements Serializable { ... } (y declarar serialVersionUID).",
                        "2. El flujo out nunca se cierra: los datos pueden no volcarse al disco y el recurso queda abierto. Corrección: usar try-with-resources:",
                        "   try (ObjectOutputStream out =",
                        "            new ObjectOutputStream(new FileOutputStream(\"p.dat\"))) {",
                        "       out.writeObject(p);",
                        "   }",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Analiza la siguiente escritura y determina qué instrucciones de lectura recuperan los datos y cuántos bytes ocupa cada valor.",
                    "puntos": "1,5 ptos",
                    "code": """DataOutputStream out =
    new DataOutputStream(new FileOutputStream("datos.dat"));
out.writeInt(7);
out.writeDouble(2.5);
out.writeUTF("Hola");
out.close();""",
                    "respuesta": [
                        "Los datos se leen en el MISMO orden en que se escribieron:",
                        "   int i = in.readInt();      // 7 (4 bytes)",
                        "   double d = in.readDouble(); // 2.5 (8 bytes)",
                        "   String s = in.readUTF();    // \"Hola\" (2 bytes de longitud + 4 de texto = 6 bytes)",
                        "Tamaño total del fichero: 4 + 8 + 6 = 18 bytes.",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 3: Ejercicios de programación",
            "puntos": "8 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "Crea la clase Persona (nombre, edad) que implemente Serializable y un programa que guarde una lista de 3 personas en personas.dat con ObjectOutputStream y después la lea con ObjectInputStream mostrándola por pantalla.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.io.*;",
                        "import java.util.*;",
                        "public class GuardarPersonas {",
                        "    static class Persona implements Serializable {",
                        "        private static final long serialVersionUID = 1L;",
                        "        String nombre; int edad;",
                        "        Persona(String nombre, int edad) { this.nombre = nombre; this.edad = edad; }",
                        "        public String toString() { return nombre + \", \" + edad + \" anios\"; }",
                        "    }",
                        "    public static void main(String[] args) {",
                        "        List<Persona> personas = List.of(",
                        "            new Persona(\"Ana\", 22),",
                        "            new Persona(\"Luis\", 19),",
                        "            new Persona(\"Eva\", 24));",
                        "        try (ObjectOutputStream out =",
                        "                new ObjectOutputStream(new FileOutputStream(\"personas.dat\"))) {",
                        "            out.writeObject(personas);",
                        "        } catch (IOException e) {",
                        "            System.out.println(\"Error al guardar: \" + e.getMessage());",
                        "        }",
                        "        try (ObjectInputStream in =",
                        "                new ObjectInputStream(new FileInputStream(\"personas.dat\"))) {",
                        "            List<Persona> leidas = (List<Persona>) in.readObject();",
                        "            for (Persona p : leidas) System.out.println(p);",
                        "        } catch (IOException | ClassNotFoundException e) {",
                        "            System.out.println(\"Error al leer: \" + e.getMessage());",
                        "        }",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Escribe un programa que copie un fichero binario (por ejemplo foto.jpg) a otro (copia_foto.jpg) leyendo byte a byte con FileInputStream y escribiendo con FileOutputStream. Muestra cuántos bytes se copiaron.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.io.*;",
                        "public class CopiarFichero {",
                        "    public static void main(String[] args) {",
                        "        File origen = new File(\"foto.jpg\");",
                        "        File destino = new File(\"copia_foto.jpg\");",
                        "        int contador = 0;",
                        "        try (FileInputStream in = new FileInputStream(origen);",
                        "             FileOutputStream out = new FileOutputStream(destino)) {",
                        "            int b;",
                        "            while ((b = in.read()) != -1) {",
                        "                out.write(b);",
                        "                contador++;",
                        "            }",
                        "        } catch (IOException e) {",
                        "            System.out.println(\"Error: \" + e.getMessage());",
                        "        }",
                        "        System.out.println(\"Copiados \" + contador + \" bytes\");",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Escribe un programa que pida 5 números double por teclado, los guarde en numeros.dat con DataOutputStream y después los lea con DataInputStream mostrando la suma y la media.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.io.*;",
                        "import java.util.Scanner;",
                        "public class DatosBinarios {",
                        "    public static void main(String[] args) throws IOException {",
                        "        Scanner sc = new Scanner(System.in);",
                        "        try (DataOutputStream out =",
                        "                new DataOutputStream(new FileOutputStream(\"numeros.dat\"))) {",
                        "            for (int i = 0; i < 5; i++) {",
                        "                System.out.print(\"Numero \" + (i + 1) + \": \");",
                        "                out.writeDouble(sc.nextDouble());",
                        "            }",
                        "        }",
                        "        double suma = 0; int contador = 0;",
                        "        try (DataInputStream in =",
                        "                new DataInputStream(new FileInputStream(\"numeros.dat\"))) {",
                        "            while (in.available() > 0) {",
                        "                suma += in.readDouble();",
                        "                contador++;",
                        "            }",
                        "        }",
                        "        System.out.println(\"Suma: \" + suma);",
                        "        System.out.println(\"Media: \" + (suma / contador));",
                        "    }",
                        "}",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Con java.nio.file: lee el fichero alumnos.csv (formato nombre;nota, una línea por alumno) con Files.readAllLines, calcula la nota media y escribe el resultado en media.txt con Files.write.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.nio.file.*;",
                        "import java.io.IOException;",
                        "import java.util.*;",
                        "public class NotaMedia {",
                        "    public static void main(String[] args) throws IOException {",
                        "        Path ruta = Paths.get(\"alumnos.csv\");",
                        "        List<String> lineas = Files.readAllLines(ruta);",
                        "        double suma = 0; int n = 0;",
                        "        for (String linea : lineas) {",
                        "            String[] campos = linea.split(\";\");",
                        "            if (campos.length == 2) {",
                        "                suma += Double.parseDouble(campos[1].trim());",
                        "                n++;",
                        "            }",
                        "        }",
                        "        double media = (n > 0) ? suma / n : 0;",
                        "        Files.write(Paths.get(\"media.txt\"),",
                        "            Arrays.asList(\"Alumnos: \" + n, \"Nota media: \" + media));",
                        "        System.out.println(\"Media calculada: \" + media);",
                        "    }",
                        "}",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 4: Pregunta teórica",
            "puntos": "1 punto",
            "preguntas": [
                {
                    "tipo": "teorica",
                    "enunciado": "Explica las diferencias entre ficheros de texto y ficheros binarios, y entre serializar objetos y escribir los datos manualmente. ¿Qué ventajas aporta java.nio.file frente a java.io.File?",
                    "respuesta": [
                        "Texto: legible por humanos, se lee con Scanner/BufferedReader; binario: compacto y sin conversión de tipos, se lee con flujos de bytes o de datos.",
                        "Serializar guarda el objeto completo (estado y estructura) con un solo writeObject/readObject, pero el formato depende de la clase y su versión (por eso existe serialVersionUID); escribir manualmente da control total del formato y es más estable entre versiones, pero hay que escribir cada campo.",
                        "java.nio.file (Path, Files) ofrece métodos de alto nivel (readAllLines, write, walk, copy, move...), trata las rutas de forma multiplataforma y permite operaciones atómicas; java.io.File es más antiguo, con menos funcionalidad y métodos que devuelven boolean sin informar del error.",
                    ],
                },
            ],
        },
    ],
})

# =====================================================================
# EXAMEN 10 — Programación funcional: lambdas, streams, Optional y java.time
# =====================================================================
EXAMENES.append({
    "titulo": "EXAMEN DE PROGRAMACIÓN – 1º GRADO SUPERIOR",
    "subtitulo": "Examen 10 · Programación funcional: lambdas, streams, Optional y java.time",
    "temas": [
        "Expresiones lambda e interfaces funcionales (Predicate, Function, Consumer, Supplier)",
        "Referencias a métodos (method references)",
        "El API Stream: filter, map, sorted, limit, reduce, collect",
        "Optional: orElse, orElseGet, isPresent, map",
        "Fecha y hora con java.time: LocalDate, LocalDateTime, Period, DateTimeFormatter",
    ],
    "puntuacion": 16,
    "info": [("Convocatoria", "Evaluación 3"), ("Valoración", "Programación funcional")],
    "secciones": [
        {
            "titulo": "Parte 1: Opción múltiple y Verdadero/Falso",
            "puntos": "4 puntos — 0,5 cada pregunta",
            "preguntas": [
                {
                    "tipo": "test",
                    "enunciado": "¿Cuál de estas expresiones es una lambda válida para sumar dos enteros?",
                    "opciones": ["a) (a, b) -> a + b", "b) a + b -> (a, b)", "c) (a, b) => a + b", "d) int a, int b -> a + b"],
                    "respuesta": "a) (a, b) -> a + b — La sintaxis de una lambda es (parámetros) -> expresión; el tipo de los parámetros se infiere.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué imprime el siguiente código?",
                    "code": """List<Integer> nums = List.of(1, 2, 3, 4, 5);
int r = nums.stream()
            .filter(n -> n % 2 == 0)
            .mapToInt(n -> n * n)
            .sum();
System.out.println(r);""",
                    "opciones": ["a) 6", "b) 20", "c) 55", "d) 30"],
                    "respuesta": "b) 20 — Filtra los pares (2 y 4), calcula sus cuadrados (4 y 16) y los suma: 4 + 16 = 20.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "Un stream solo puede recorrerse una vez: después de consumirlo no puede reutilizarse.",
                    "respuesta": "Verdadero — Un stream es un flujo de una sola pasada; intentar volver a operar sobre el mismo stream lanza IllegalStateException.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué imprime este código?",
                    "code": """Optional<String> o = Optional.empty();
System.out.println(o.orElse("vacio"));""",
                    "opciones": ["a) Optional.empty", "b) vacio", "c) null", "d) NoSuchElementException"],
                    "respuesta": "b) vacio — orElse() devuelve el valor si está presente y, si no, el valor alternativo indicado.",
                },
                {
                    "tipo": "test",
                    "enunciado": "La referencia a método String::length equivale a la lambda de una expresión. ¿Cuál?",
                    "opciones": ["a) () -> String.length()", "b) s -> s.length()", "c) s -> String.length(s)", "d) String -> length"],
                    "respuesta": "b) s -> s.length() — Una method reference sobre un parámetro de instancia equivale a invocar el método sobre ese parámetro.",
                },
                {
                    "tipo": "vf",
                    "enunciado": "Collectors.toList() recoge los elementos de un stream en una nueva lista.",
                    "respuesta": "Verdadero — Es el collector más habitual: stream.collect(Collectors.toList()) devuelve una List con los elementos resultantes.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué fecha imprime este código?",
                    "code": """LocalDate hoy = LocalDate.of(2026, 3, 10);
System.out.println(hoy.plusDays(20));""",
                    "opciones": ["a) 2026-03-30", "b) 2026-03-20", "c) 2026-04-01", "d) 2026-04-30"],
                    "respuesta": "a) 2026-03-30 — LocalDate es inmutable; plusDays(20) devuelve una nueva fecha: 10 de marzo + 20 días = 30 de marzo.",
                },
                {
                    "tipo": "test",
                    "enunciado": "¿Qué interfaz funcional se usa para representar una condición que devuelve true o false sobre un objeto?",
                    "opciones": ["a) Consumer<T>", "b) Supplier<T>", "c) Predicate<T>", "d) Function<T, R>"],
                    "respuesta": "c) Predicate<T> — Su método test(T) devuelve boolean; es la interfaz típica para filter().",
                },
            ],
        },
        {
            "titulo": "Parte 2: Análisis y depuración",
            "puntos": "3 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "El siguiente programa tiene 2 errores. Encuéntralos, explica por qué fallan y propón la corrección.",
                    "puntos": "1,5 ptos",
                    "code": """import java.util.*;
import java.util.stream.*;
public class StreamBug {
    public static void main(String[] args) {
        List<String> nombres = List.of("Ana", "Bea", "Carlos");
        Stream<String> st = nombres.stream();
        st.forEach(System.out::println);
        st.filter(s -> s.length() > 2).forEach(System.out::println);

        int limite = 2;
        nombres.stream().filter(s -> s.length() >= limite);
        limite = 3;
    }
}""",
                    "respuesta": [
                        "1. Se usa el mismo stream st dos veces: tras forEach() el stream está consumido y la segunda operación lanza IllegalStateException (stream has already been operated upon). Corrección: crear un stream nuevo por operación: nombres.stream().filter(...)...",
                        "2. La variable limite se usa dentro de una lambda y después se reasigna (limite = 3): las lambdas solo pueden capturar variables effectively final → error de compilación. Corrección: no reasignarla o usar otra variable.",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Analiza el siguiente pipeline y determina qué nombres imprime y en qué orden (traza manual).",
                    "puntos": "1,5 ptos",
                    "code": """List<String> nombres = List.of("Ana", "Bea", "Carlos", "David", "Eva");
nombres.stream()
       .filter(s -> s.length() >= 3)
       .sorted()
       .forEach(System.out::println);""",
                    "respuesta": [
                        "1. filter(s -> s.length() >= 3): todas las cadenas tienen 3 o más letras, así que ninguna se descarta (Eva tiene 3).",
                        "2. sorted() ordena alfabéticamente (orden natural de String).",
                        "Salida: Ana, Bea, Carlos, David, Eva.",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 3: Ejercicios de programación",
            "puntos": "8 puntos",
            "preguntas": [
                {
                    "tipo": "ejercicio",
                    "enunciado": "Con streams: dada la lista [1, 2, 3, 4, 5, 6, 7, 8], calcula la suma de los cuadrados de los números pares y muestra los 3 mayores pares en orden descendente.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.util.*;",
                        "import java.util.stream.*;",
                        "public class StreamsNumeros {",
                        "    public static void main(String[] args) {",
                        "        List<Integer> numeros = List.of(1, 2, 3, 4, 5, 6, 7, 8);",
                        "        int sumaCuadradosPares = numeros.stream()",
                        "                .filter(n -> n % 2 == 0)",
                        "                .mapToInt(n -> n * n)",
                        "                .sum();",
                        "        System.out.println(\"Suma de cuadrados de pares: \" + sumaCuadradosPares);",
                        "        List<Integer> paresDesc = numeros.stream()",
                        "                .filter(n -> n % 2 == 0)",
                        "                .sorted(Comparator.reverseOrder())",
                        "                .limit(3)",
                        "                .collect(Collectors.toList());",
                        "        System.out.println(\"Tres mayores pares: \" + paresDesc);",
                        "    }",
                        "}",
                        "Salida: Suma de cuadrados de pares: 120 (4+16+36+64) y Tres mayores pares: [8, 6, 4].",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Con una lista de alumnos (nombre, nota), usa streams para: (a) calcular la nota media, (b) mostrar los aprobados y (c) mostrar el alumno con mejor nota.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.util.*;",
                        "public class StreamsAlumnos {",
                        "    static class Alumno {",
                        "        String nombre; double nota;",
                        "        Alumno(String nombre, double nota) { this.nombre = nombre; this.nota = nota; }",
                        "        public String toString() { return nombre; }",
                        "    }",
                        "    public static void main(String[] args) {",
                        "        List<Alumno> alumnos = List.of(",
                        "            new Alumno(\"Ana\", 7.5), new Alumno(\"Luis\", 4.9),",
                        "            new Alumno(\"Eva\", 9.2), new Alumno(\"Juan\", 5.5));",
                        "        double media = alumnos.stream()",
                        "                .mapToDouble(a -> a.nota)",
                        "                .average()",
                        "                .orElse(0);",
                        "        System.out.println(\"Nota media: \" + media);",
                        "        System.out.println(\"Aprobados:\");",
                        "        alumnos.stream()",
                        "                .filter(a -> a.nota >= 5)",
                        "                .forEach(System.out::println);",
                        "        Alumno mejor = alumnos.stream()",
                        "                .max(Comparator.comparingDouble(a -> a.nota))",
                        "                .orElse(null);",
                        "        System.out.println(\"Mejor: \" + mejor);",
                        "    }",
                        "}",
                        "Salida: Nota media: 6.775; aprobados Ana, Eva y Juan; Mejor: Eva.",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Implementa un método que devuelva Optional<Alumno> al buscar un alumno por nombre en una lista (sin distinguir mayúsculas). En el main usa orElse() y ifPresentOrElse() para gestionar la ausencia.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.util.*;",
                        "public class BuscarAlumno {",
                        "    static class Alumno {",
                        "        String nombre; double nota;",
                        "        Alumno(String nombre, double nota) { this.nombre = nombre; this.nota = nota; }",
                        "        public String toString() { return nombre + \" (\" + nota + \")\"; }",
                        "    }",
                        "    static Optional<Alumno> buscar(List<Alumno> lista, String nombre) {",
                        "        for (Alumno a : lista) {",
                        "            if (a.nombre.equalsIgnoreCase(nombre)) return Optional.of(a);",
                        "        }",
                        "        return Optional.empty();",
                        "    }",
                        "    public static void main(String[] args) {",
                        "        List<Alumno> alumnos = List.of(",
                        "            new Alumno(\"Ana\", 7.5), new Alumno(\"Eva\", 9.2));",
                        "        Alumno encontrado = buscar(alumnos, \"eva\")",
                        "                .orElse(new Alumno(\"Desconocido\", 0));",
                        "        System.out.println(\"Resultado: \" + encontrado);",
                        "        buscar(alumnos, \"Pedro\").ifPresentOrElse(",
                        "            a -> System.out.println(\"Nota de \" + a.nombre + \": \" + a.nota),",
                        "            () -> System.out.println(\"Pedro no esta en la lista\"));",
                        "    }",
                        "}",
                        "Salida: Resultado: Eva (9.2) y Pedro no esta en la lista.",
                    ],
                },
                {
                    "tipo": "ejercicio",
                    "enunciado": "Con java.time: pide una fecha de nacimiento (AAAA-MM-DD), calcula la edad con Period.between, los días vividos con ChronoUnit.DAYS.between y muestra la fecha de hoy formateada como dd/MM/yyyy.",
                    "puntos": "2 ptos",
                    "respuesta": [
                        "import java.time.*;",
                        "import java.time.format.DateTimeFormatter;",
                        "import java.util.Scanner;",
                        "public class Fechas {",
                        "    public static void main(String[] args) {",
                        "        Scanner sc = new Scanner(System.in);",
                        "        System.out.print(\"Fecha de nacimiento (AAAA-MM-DD): \");",
                        "        LocalDate nacimiento = LocalDate.parse(sc.nextLine());",
                        "        LocalDate hoy = LocalDate.now();",
                        "        Period periodo = Period.between(nacimiento, hoy);",
                        "        System.out.println(\"Edad: \" + periodo.getYears() + \" anios, \"",
                        "            + periodo.getMonths() + \" meses y \" + periodo.getDays() + \" dias\");",
                        "        long diasVividos = ChronoUnit.DAYS.between(nacimiento, hoy);",
                        "        System.out.println(\"Dias vividos: \" + diasVividos);",
                        "        DateTimeFormatter formato = DateTimeFormatter.ofPattern(\"dd/MM/yyyy\");",
                        "        System.out.println(\"Hoy: \" + hoy.format(formato));",
                        "    }",
                        "}",
                    ],
                },
            ],
        },
        {
            "titulo": "Parte 4: Pregunta teórica",
            "puntos": "1 punto",
            "preguntas": [
                {
                    "tipo": "teorica",
                    "enunciado": "Explica qué es una expresión lambda y qué es una interfaz funcional. ¿Qué ventajas aportan los streams frente a los bucles tradicionales? Pon un ejemplo comparativo.",
                    "respuesta": [
                        "Una lambda es una función anónima: (parámetros) -> expresión, que puede asignarse a una interfaz funcional. Una interfaz funcional tiene un único método abstracto (Predicate, Function, Consumer, Supplier, Comparator...).",
                        "Ventajas de los streams: código más conciso y legible, operaciones encadenadas (filter, map, sorted, limit...), evaluación perezosa (las operaciones intermedias solo se ejecutan si hay una operación final) y facilidad para paralelizar con parallelStream().",
                        "Ejemplo: filtrar y mostrar nombres que empiezan por 'A'.",
                        "Bucle tradicional:",
                        "   for (String s : nombres) if (s.startsWith(\"A\")) System.out.println(s);",
                        "Con streams:",
                        "   nombres.stream().filter(s -> s.startsWith(\"A\")).forEach(System.out::println);",
                    ],
                },
            ],
        },
    ],
})


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    out_dir = os.path.dirname(os.path.abspath(__file__))
    args = sys.argv[1:] if len(sys.argv) > 1 else ["docx", "pdf"]
    want_docx = "docx" in args
    want_pdf = "pdf" in args
    dark = "dark" in args or "oscuro" in args

    print("Fuente Unicode disponible:", bool(find_font("arial.ttf")))
    print("Modo PDF:", "OSCURO" if dark else "claro")

    for i, ex in enumerate(EXAMENES, start=6):
        nombre = f"Examen_{i}_Java_{ex['subtitulo'].split('·')[-1].strip().replace(' ','_')}"
        # Windows no permite estos caracteres en nombres de archivo
        for c in ':;*?"<>|':
            nombre = nombre.replace(c, "-")
        nombre = nombre.replace("/", "-").replace("\\", "-")[:80]
        base = os.path.join(out_dir, nombre)
        print(f"\n=== Examen {i}: {ex['subtitulo']} ===")

        if want_docx:
            path = base + ".docx"
            try:
                render_docx(ex, path, dark=dark)
                print("   DOCX OK:", os.path.basename(path))
            except Exception as e:
                print("   DOCX ERROR:", e)

        if want_pdf:
            path = base + ".pdf"
            try:
                PDFExam(ex, dark=dark).build(path)
                print("   PDF  OK:", os.path.basename(path))
            except Exception as e:
                import traceback
                print("   PDF  ERROR:", e)
                traceback.print_exc()

    print("\nFinalizado.")


if __name__ == "__main__":
    main()
