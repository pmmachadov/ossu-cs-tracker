#!/usr/bin/env python
"""Genera un documento Word con el examen estilo IOC."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ─────────────────────────────── ESTILOS ───────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

# ─────────────────────────── PORTADA / ENCABEZADO ──────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('EXAMEN DE PROGRAMACIÓN – 1º DAW')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Estilo IOC (Institut Obert de Catalunya)')
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Basado en los contenidos del repositorio:\n'
                  'Retos IOC 1-6 + Curso YouTube "Aula en la nube" – MEGA Curso JAVA [DAM-DAW]')
run3.font.size = Pt(10)
run3.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()  # espacio

# ──────────── TABLA INFORMATIVA ────────────
table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light Shading Accent 1'

info = [
    ('Asignatura', 'Programación (PRG)'),
    ('Duración', '2 horas'),
    ('Curso', '1º Desarrollo de Aplicaciones Web (DAW)'),
    ('Puntuación total', '16 puntos (mínimo 8 para aprobar)'),
]
for i, (k, v) in enumerate(info):
    row = table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.style.font.size = Pt(11)

doc.add_paragraph()

# ──────────────────────────────────────────────────────────────────────────
# FUNCIÓN AUXILIAR
# ──────────────────────────────────────────────────────────────────────────
def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    return h

def add_question(number, text, points=''):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{number}. ')
    run_num.bold = True
    run_num.font.size = Pt(11.5)
    p.add_run(text).font.size = Pt(11.5)
    if points:
        run_pts = p.add_run(f'  ({points})')
        run_pts.font.size = Pt(10)
        run_pts.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run_pts.italic = True
    return p

def add_code(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_options(options_list):
    for letter, text in options_list:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(f'{letter}) {text}')
        run.font.size = Pt(11)
    return p

# =========================================================================
# PARTE 1
# =========================================================================
add_heading_custom('Parte 1: Opción múltiple y Verdadero/Falso', level=1)
p = doc.add_paragraph('(4 puntos — 0.5 cada pregunta)')
p.runs[0].italic = True

# 1.1
add_question('1.1', 'Dado el código:')
add_code('''int a = 5;
double b = 2.0;
System.out.println(a / b);''')
add_options([('a', '2'), ('b', '2.0'), ('c', '2.5'), ('d', 'Error de compilación')])

# 1.2
add_question('1.2', '¿Cuál de estos bucles ejecuta el cuerpo al menos una vez?')
add_options([('a', 'while (cond) { }'), ('b', 'for (int i=0; i<n; i++) { }'),
             ('c', 'do { } while (cond);'), ('d', 'Todos ejecutan el cuerpo al menos una vez')])

# 1.3
add_question('1.3', 'Indica si es Verdadero o Falso:')
add_code('''String a = "Hola";
String b = "Hola";
System.out.println(a == b);''')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.5)
p.add_run('Imprime true siempre en cualquier ejecución.').font.size = Pt(11)

# 1.4
add_question('1.4', 'Dado:')
add_code('''int[] nums = {3, 7, 2, 9, 5};
Arrays.sort(nums);
System.out.println(Arrays.binarySearch(nums, 7));''')
add_options([('a', '3'), ('b', '4'), ('c', '7'), ('d', '1')])

# 1.5
add_question('1.5', '¿Cuál de estas NO es una convención válida de nombre de variable en Java?')
add_options([('a', '$importe'), ('b', '_contador'), ('c', '2daOpcion'), ('d', 'totalIVA')])

# 1.6
add_question('1.6', 'V/F: En Java, cuando se pasa un array a un método, se pasa por valor la referencia, por lo que modificar elementos dentro del método modifica el array original.')

# 1.7
add_question('1.7', '¿Qué permite la palabra clave abstract en una clase?')
add_options([('a', 'Crear objetos directamente de esa clase'),
             ('b', 'Definir métodos sin implementación que las subclases deben implementar'),
             ('c', 'Acelerar la compilación'),
             ('d', 'Hacer que todos sus métodos sean private')])

# 1.8
add_question('1.8', 'Dado:')
add_code('''Random r = new Random(42);
System.out.println(r.nextInt(10));''')
add_options([('a', 'Cada ejecución imprime un número diferente'),
             ('b', 'Siempre imprime el mismo número porque la semilla es fija'),
             ('c', 'Lanza excepción porque falta import java.util.Random'),
             ('d', 'Imprime 42')])

doc.add_page_break()

# =========================================================================
# PARTE 2
# =========================================================================
add_heading_custom('Parte 2: Análisis y depuración', level=1)
p = doc.add_paragraph('(3 puntos)')
p.runs[0].italic = True

# 2.1
add_question('2.1', 'El siguiente código tiene 3 errores. Encuéntralos y explica por qué fallan:', '1.5 ptos')
add_code('''public class GestionNotas {
    public static void main(String[] args) {
        float[] notas;
        notas = {4.5f, 7.2f, 3.8f, 9.1f, 6.0f};
        
        for (int i = 0; i <= notas.length; i++) {
            if (notas[i] >= 5)
                System.out.println("Aprobado");
            else
                System.out.println("Suspenso");
        }
        
        System.out.println("Media: " + calcularMedia(notas[]));
    }
    
    static float calcularMedia(float[] n) {
        float suma = 0;
        for (int i = 0; i < n.length; i++)
            suma = suma + n[i];
        return suma;
    }
}''')

# 2.2
add_question('2.2', 'Analiza y di qué imprime este programa:', '1.5 ptos')
add_code('''public class Misterio {
    public static void main(String[] args) {
        int x = 10;
        int[] arr = {3, 7, 2, 8, 1};
        
        for (int i = 0; i < arr.length - 1; i++) {
            for (int j = i + 1; j < arr.length; j++) {
                if (arr[i] > arr[j]) {
                    int temp = arr[i];
                    arr[i] = arr[j];
                    arr[j] = temp;
                }
            }
        }
        
        System.out.println(Arrays.toString(arr));
        System.out.println(buscar(arr, 7));
    }
    
    static int buscar(int[] a, int valor) {
        int izq = 0, der = a.length - 1;
        while (izq <= der) {
            int medio = (izq + der) / 2;
            if (a[medio] == valor) return medio;
            if (a[medio] < valor) izq = medio + 1;
            else der = medio - 1;
        }
        return -1;
    }
}''')

doc.add_page_break()

# =========================================================================
# PARTE 3
# =========================================================================
add_heading_custom('Parte 3: Ejercicios de programación', level=1)
p = doc.add_paragraph('(8 puntos)')
p.runs[0].italic = True

# 3.1
add_question('3.1', 'Arrays y métodos', '2 ptos')
p = doc.add_paragraph()
p.add_run('Escribe un programa completo que:').font.size = Pt(11)
items_31 = [
    'Pida por teclado 8 números enteros y los guarde en un array.',
    'Implemente un método static int contarPares(int[] arr) que devuelva cuántos números pares contiene.',
    'Implemente un método static int[] invertir(int[] arr) que devuelva un nuevo array con los elementos en orden inverso.',
    'Muestre el resultado de ambos métodos.',
]
for item in items_31:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(11)

# 3.2
add_question('3.2', 'Programación Orientada a Objetos', '2.5 ptos')
p = doc.add_paragraph()
p.add_run('Diseña las clases para gestionar una Biblioteca digital:').font.size = Pt(11)
items_32 = [
    'Clase Libro: atributos privados titulo (String), autor (String), anyoPublicacion (int), prestado (boolean). Constructor con todos los parámetros. Métodos getters y setters. Método prestar() que marque como prestado si no lo está. Método devolver() que marque como disponible. Método toString() que muestre la información del libro.',
    'Clase Biblioteca: atributo libros (ArrayList<Libro>). Método agregarLibro(Libro l). Método listarDisponibles() que muestre solo los libros no prestados. Método buscarPorAutor(String autor) que devuelva los libros de ese autor.',
]
for item in items_32:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(11)

# 3.3
add_question('3.3', 'Ficheros', '2 ptos')
p = doc.add_paragraph()
p.add_run('Escribe un programa que:').font.size = Pt(11)
items_33 = [
    'Lea un archivo de texto llamado alumnos.txt donde cada línea tiene el formato: nombre nota1 nota2 nota3. Ejemplo: Ana 7.5 8.0 6.5',
    'Para cada alumno, calcule la media de sus notas.',
    'Escriba en un nuevo archivo resultados.txt el nombre y la media, con el formato: Ana: 7.33',
    'Gestione correctamente las excepciones con try-with-resources.',
]
for item in items_33:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(11)

# 3.4
add_question('3.4', 'Herencia y polimorfismo', '1.5 ptos')
p = doc.add_paragraph()
p.add_run('Implementa:').font.size = Pt(11)
items_34 = [
    'Una clase abstracta Figura con un método abstracto double area() y un método void mostrar() que imprima el área.',
    'Una clase Circulo que herede de Figura con atributo radio.',
    'Una clase Rectangulo que herede de Figura con atributos ancho y alto.',
    'Un programa principal que cree un array de 3 figuras (mezcla de círculos y rectángulos) y muestre sus áreas usando polimorfismo.',
]
for item in items_34:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(11)

# =========================================================================
# PARTE 4
# =========================================================================
add_heading_custom('Parte 4: Pregunta teórica', level=1)
p = doc.add_paragraph('(1 punto)')
p.runs[0].italic = True

add_question('4.1', 'Explica la diferencia entre:', '1 pto')
items_41 = [
    'while y do-while',
    'ArrayList y array tradicional (int[])',
    '== y .equals() al comparar Strings',
    'public, private y protected',
]
for item in items_41:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(11)

# =========================================================================
# ESPACIO PARA RESPUESTAS
# =========================================================================
doc.add_page_break()
add_heading_custom('Espacio para respuestas', level=1)
p = doc.add_paragraph()
p.add_run('(Páginas adicionales para desarrollar las soluciones)').italic = True

for i in range(6):
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

# ──────────────────────────── GUARDAR ────────────────────────────
output_path = r'C:\Users\Pablo\Desktop\Pablo\anki-cards-completo\Examen_Programacion_DAW_Estilo_IOC.docx'
doc.save(output_path)
print(f'✅ Documento guardado en:\n{output_path}')
