#!/usr/bin/env python3
"""Generate DOCX by parsing the existing HTML file. Font size 12."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, re

BASE = r"C:\Users\Pablo\Desktop\Pablo\anki-cards-completo\pruebas\programacion\JAVA IOC\JAVA YOUTUBE\05-POO_Programacion_Orientada_Objetos"
HTML_PATH = os.path.join(BASE, "examen", "examen-poo.html")
OUTPUT_PATH = os.path.join(BASE, "examen", "Examen_POO_Java_Tema5.docx")

with open(HTML_PATH, 'r', encoding='utf-8') as f:
    html = f.read()

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)

def add_code_run(para, text):
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    return run

def add_normal_run(para, text, bold=False, italic=False, color=None, size=None):
    run = para.add_run(text)
    run.font.name = 'Calibri'
    if bold: run.bold = True
    if italic: run.italic = True
    if color: run.font.color.rgb = RGBColor(*color)
    if size: run.font.size = Pt(size)
    return run

def strip_html_tags(text):
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'</?code>', '', text)
    text = re.sub(r'</?strong>', '', text)
    text = re.sub(r'</?em>', '', text)
    text = re.sub(r'</?ol>', '', text)
    text = re.sub(r'</?ul>', '', text)
    text = re.sub(r'</?li>', '\n  - ', text)
    text = re.sub(r'</?p>', '\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def extract_pre_content(text):
    m = re.search(r'<pre>(.*?)</pre>', text, re.DOTALL)
    return m.group(1).strip() if m else ''

def remove_pre_blocks(text):
    return re.sub(r'<pre>.*?</pre>', '', text, flags=re.DOTALL)

# Portada
title = doc.add_heading('EXAMEN DE PROGRAMACION', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('Programacion Orientada a Objetos en Java', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Tema 5: POO  |  Curso Java desde 0  |  DAM / DAW')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub2.add_run('27 videos - 56 ejercicios practicos - Incluye soluciones completas')
r2.font.size = Pt(11)
r2.font.italic = True

doc.add_page_break()

# Parsear HTML
video_groups = html.split('<div class="video-group">')
first_vg = True

for vg_html in video_groups:
    if not vg_html.strip():
        continue
    if first_vg:
        first_vg = False

    vh_match = re.search(r'<div class="video-header">\s*<h2>(.*?)</h2>\s*<div class="sub">(.*?)</div>', vg_html, re.DOTALL)
    if not vh_match:
        continue

    video_title = strip_html_tags(vh_match.group(1)).strip()
    video_sub = strip_html_tags(vh_match.group(2)).strip()

    heading = doc.add_heading(video_title, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    pPr = heading._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '1A3A5C')
    shading.set(qn('w:val'), 'clear')
    pPr.append(shading)

    sub_p = doc.add_paragraph()
    add_normal_run(sub_p, video_sub, italic=True, size=10)

    exercises = vg_html.split('<div class="exercise">')
    for ex_html in exercises[1:]:
        q_match = re.search(r'<div class="question">(.*?)</div>\s*<div class="solution">', ex_html, re.DOTALL)
        if not q_match:
            continue
        q_html = q_match.group(1)

        s_match = re.search(r'<div class="solution">(.*?)</div>', ex_html, re.DOTALL)
        s_html = s_match.group(1) if s_match else ''

        q_heading = ''
        q_title_match = re.search(r'<h3>(.*?)</h3>', q_html, re.DOTALL)
        if q_title_match:
            q_heading = strip_html_tags(q_title_match.group(1))

        q_body = re.sub(r'<h3>.*?</h3>', '', q_html, flags=re.DOTALL)
        q_pre = extract_pre_content(q_body)
        q_text = strip_html_tags(remove_pre_blocks(q_body))

        s_pre = extract_pre_content(s_html)
        s_text = strip_html_tags(remove_pre_blocks(s_html))

        # Ejercicio
        ex_heading = doc.add_heading(q_heading, level=2)
        for run in ex_heading.runs:
            run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

        # Enunciado
        p = doc.add_paragraph()
        add_normal_run(p, 'ENUNCIADO:', bold=True, color=(0xcc, 0x33, 0x33))
        if q_text.strip():
            p = doc.add_paragraph(q_text.strip())
        if q_pre:
            p = doc.add_paragraph()
            add_code_run(p, q_pre)

        # Solucion
        p = doc.add_paragraph()
        add_normal_run(p, 'SOLUCION:', bold=True, color=(0x28, 0xa7, 0x45))
        if s_text.strip():
            p = doc.add_paragraph(s_text.strip())
        if s_pre:
            p = doc.add_paragraph()
            add_code_run(p, s_pre)

        doc.add_paragraph('-' * 60)

doc.save(OUTPUT_PATH)
print(f'DOCX creado: {OUTPUT_PATH}')
size = os.path.getsize(OUTPUT_PATH)
print(f'Tamano: {size} bytes ({size/1024:.1f} KB)')

ej = html.count('<div class="question">')
sol = html.count('<div class="solution">')
print(f'Ejercicios en HTML: {ej}')
print(f'Soluciones en HTML: {sol}')
