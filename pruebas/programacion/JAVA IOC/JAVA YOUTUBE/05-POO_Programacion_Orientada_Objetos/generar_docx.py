#!/usr/bin/env python3
"""Generate the exam as a Word DOCX document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

OUTPUT_DIR = r"C:\Users\Pablo\Desktop\Pablo\anki-cards-completo\pruebas\programacion\JAVA IOC\JAVA YOUTUBE\05-POO_Programacion_Orientada_Objetos"
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "Examen_POO_Java_Tema5.docx")

doc = Document()

# Configure styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

# Title
title = doc.add_heading('EXAMEN DE PROGRAMACIÓN', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('Programación Orientada a Objetos en Java', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('TEMA 5 \u00b7 POO  \u2014  Curso Java desde 0  \u2014  DAM / DAW')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)

doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub2.add_run('27 videos  \u00b7  54 ejercicios pr\u00e1cticos  \u00b7  Incluye soluciones completas')
r2.font.size = Pt(10)
r2.font.italic = True

doc.add_page_break()

# ============================================================
# Define all exercises data
# ============================================================

exercises = [
    # (video_num, video_title, subtitle, exercise_num, question_text, solution_text)
    
    # V-01
    ("01", "Introducción a la POO", "Conceptos: objeto, clase, instancia, estado, comportamiento, mensajes",
     "1.1", "Identificar elementos POO",
     "Dado el siguiente escenario: \"Tenemos una aplicaci\u00f3n para gestionar una biblioteca. Necesitamos representar libros, cada uno con su t\u00edtulo, autor, ISBN, a\u00f1o de publicaci\u00f3n, y si est\u00e1 prestado o no. Los libros se pueden prestar, devolver y consultar.\"\n\nIdentifica y explica en el contexto de este escenario:\n1. \u00bfCu\u00e1l ser\u00eda la clase?\n2. \u00bfQu\u00e9 son los atributos?\n3. \u00bfQu\u00e9 son los m\u00e9todos?\n4. \u00bfQu\u00e9 ser\u00eda una instancia? Pon un ejemplo concreto.\n5. \u00bfQu\u00e9 es el estado de un objeto? Pon un ejemplo.\n6. \u00bfQu\u00e9 es un mensaje entre objetos? Pon un ejemplo.",
     "1. Clase: Libro - el molde o plantilla para crear objetos libro.\n2. Atributos: titulo, autor, isbn, anioPublicacion, prestado (boolean). Definen el estado de cada libro.\n3. M\u00e9todos: prestar(), devolver(), consultarDatos(). Definen el comportamiento del libro.\n4. Instancia: Un libro concreto. Ej: Libro miLibro = new Libro(\"Cien A\u00f1os de Soledad\", \"Garc\u00eda M\u00e1rquez\", \"978-84-376-0494-7\", 1967);\n5. Estado: Valor actual de todos los atributos. Ej: {titulo=\"Cien A\u00f1os...\", autor=\"Garc\u00eda M\u00e1rquez\", prestado=true}\n6. Mensaje: Llamada de un objeto a un m\u00e9todo de otro. Ej: biblioteca.prestarLibro(miLibro)"),

    ("01", "Introducción a la POO", "Conceptos: objeto, clase, instancia, estado, comportamiento, mensajes",
     "1.2", "Crisis del software y convenciones",
     "Responde brevemente:\n1. \u00bfQu\u00e9 problemas surgieron en la llamada \"crisis del software\" de los a\u00f1os 60-70 que llevaron al desarrollo de la POO?\n2. Explica las convenciones de nomenclatura en Java para: clases, variables y m\u00e9todos.\n3. \u00bfPor qu\u00e9 decimos que una clase es un \"molde\" o \"plantilla\"? \u00bfQu\u00e9 relaci\u00f3n tiene con el concepto de \"instancia\"?",
     "1. Crisis del software: Los programas crec\u00edan y el paradigma orientado a procedimientos generaba: c\u00f3digo dif\u00edcil de mantener, c\u00f3digo poco reutilizable, proyectos fuera de plazo y presupuesto, y software de baja calidad. La POO surgi\u00f3 como soluci\u00f3n.\n2. Clases: PascalCase (Persona, CuentaBancaria). Variables: camelCase (nombrePersona). M\u00e9todos: camelCase + par\u00e9ntesis (obtenerDatos()).\n3. Una clase define la estructura y comportamiento que tendr\u00e1n todos los objetos de ese tipo. Una instancia es un objeto concreto creado a partir de ese molde, con valores espec\u00edficos."),

    # V-02
    ("02", "Elementos básicos de la POO", "Clase, constructor, m\u00e9todos de instancia, crear instancias con new",
     "2.1", "Dise\u00f1ar una clase Coche",
     "Dise\u00f1a una clase Coche en Java con:\n- Atributos: marca (String), modelo (String), velocidad (int, velocidad actual en km/h).\n- Un constructor con 3 par\u00e1metros que inicialice todos los atributos.\n- Un m\u00e9todo de instancia acelerar(int incremento) que aumente la velocidad.\n- Un m\u00e9todo de instancia frenar(int decremento) que disminuya la velocidad.\n- Un m\u00e9todo mostrarDatos() que imprima los datos del coche.\nEscribe el c\u00f3digo Java de la clase completa.",
     "class Coche {\n    String marca;\n    String modelo;\n    int velocidad;\n\n    Coche(String marca, String modelo, int velocidad) {\n        this.marca = marca;\n        this.modelo = modelo;\n        this.velocidad = velocidad;\n    }\n\n    void acelerar(int incremento) {\n        velocidad += incremento;\n    }\n\n    void frenar(int decremento) {\n        velocidad -= decremento;\n        if (velocidad < 0) velocidad = 0;\n    }\n\n    void mostrarDatos() {\n        System.out.println(\"Marca: \" + marca);\n        System.out.println(\"Modelo: \" + modelo);\n        System.out.println(\"Velocidad: \" + velocidad + \" km/h\");\n    }\n}"),

    ("02", "Elementos básicos de la POO", "Clase, constructor, m\u00e9todos de instancia, crear instancias con new",
     "2.2", "UML y creaci\u00f3n de objetos",
     "Dada la clase Alumno (con nombre, edad, curso, constructor y mostrarDatos):\n1. Describe el diagrama UML de esta clase (tres secciones).\n2. Escribe el c\u00f3digo para crear dos instancias de Alumno.\n3. Explica qu\u00e9 diferencia hay entre un m\u00e9todo de instancia (sin static) y un m\u00e9todo de clase (con static).",
     "1. UML:\n+------------------+\n|     Alumno       |\n+------------------+\n| - nombre: String |\n| - edad: int      |\n| - curso: String  |\n+------------------+\n| + mostrarDatos() |\n+------------------+\n\n2. Alumno a1 = new Alumno(\"Ana Garc\u00eda\", 20, \"1\u00ba DAM\");\n   Alumno a2 = new Alumno(\"Carlos L\u00f3pez\", 22, \"2\u00ba DAW\");\n\n3. M\u00e9todo de instancia (sin static): se llama sobre un objeto (a1.mostrarDatos()). Trabaja con los atributos de esa instancia. M\u00e9todo de clase (static): se llama a trav\u00e9s de la clase (Alumno.metodo()). No necesita instancia. No accede a atributos de instancia."),
]

def add_exercise(doc, v_num, v_title, v_sub, e_num, e_title, q_text, s_text):
    """Add an exercise with question and solution to the document."""
    # Video header (only changes when v_num changes)
    # We handle video headers separately in the main loop
    
    # Exercise heading
    heading = doc.add_heading(f'Ejercicio {e_num} \u2014 {e_title}', level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
    
    # Question
    q_para = doc.add_paragraph()
    q_run = q_para.add_run('✏️ ENUNCIADO:')
    q_run.bold = True
    q_run.font.color.rgb = RGBColor(0xcc, 0x33, 0x33)
    
    q_text_para = doc.add_paragraph(q_text)
    q_text_para.paragraph_format.space_after = Pt(6)
    
    # Solution
    s_para = doc.add_paragraph()
    s_run = s_para.add_run('✅ SOLUCIÓN:')
    s_run.bold = True
    s_run.font.color.rgb = RGBColor(0x28, 0xa7, 0x45)
    
    s_text_para = doc.add_paragraph(s_text)
    s_text_para.paragraph_format.space_after = Pt(12)
    
    # Separator
    doc.add_paragraph('─' * 60)

# Build document
current_video = None
for v_num, v_title, v_sub, e_num, e_title, q_text, s_text in exercises:
    if current_video != v_num:
        # New video section
        if current_video is not None:
            doc.add_page_break()
        current_video = v_num
        
        # Video header
        v_heading = doc.add_heading(f'V-{v_num}: {v_title}', level=1)
        for run in v_heading.runs:
            run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            run.font.highlight_color = None
        # Set background via XML
        from docx.oxml import OxmlElement
        for paragraph in doc.paragraphs:
            if paragraph.text.startswith(f'V-{v_num}'):
                pPr = paragraph._p.get_or_add_pPr()
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), '1A3A5C')
                shading.set(qn('w:val'), 'clear')
                pPr.append(shading)
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                    run.font.size = Pt(14)
        
        sub_p = doc.add_paragraph(v_sub)
        sub_p.runs[0].font.size = Pt(9)
        sub_p.runs[0].font.italic = True
    
    add_exercise(doc, v_num, v_title, v_sub, e_num, e_title, q_text, s_text)

# Save
doc.save(OUTPUT_PATH)
print(f"DOCX created: {OUTPUT_PATH}")
print(f"Size: {os.path.getsize(OUTPUT_PATH) if os.path.exists(OUTPUT_PATH) else 0} bytes")
