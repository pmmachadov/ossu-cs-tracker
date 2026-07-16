#!/usr/bin/env python3
"""Genera un Word con el resumen del examen, cada ejercicio en una hoja max."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import re

doc = Document()

# ── Configuración de estilos ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

# Márgenes estrechos para que quepa más
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Estilo para código
code_style = doc.styles.add_style('Codigo', 1)  # 1 = paragraph
code_font = code_style.font
code_font.name = 'Consolas'
code_font.size = Pt(9.5)
code_font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.left_indent = Cm(0.5)

def add_title(text, level=0):
    """Añade título con formato."""
    if level == 0:
        p = doc.add_heading(text, level=1)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
    elif level == 1:
        p = doc.add_heading(text, level=2)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
    elif level == 2:
        p = doc.add_heading(text, level=3)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x24, 0x24, 0x24)

def add_code(text):
    """Añade bloque de código."""
    for line in text.strip().split('\n'):
        doc.add_paragraph(line, style='Codigo')

def add_normal(text):
    """Añade párrafo normal."""
    doc.add_paragraph(text)

def add_bullet(text):
    """Añade viñeta."""
    p = doc.add_paragraph(text, style='List Bullet')

# ══════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('EXAMEN TEST 1ª EVALUACIÓN')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('TEMA 4 - ARRAYS · DAM/DAW')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('10 preguntas tipo test · 30 minutos · 2.5 puntos')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x58, 0x58, 0x58)

doc.add_paragraph()
doc.add_paragraph()
extra = doc.add_paragraph()
extra.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = extra.add_run('Basado en el vídeo: "4-26 DAM - DAW: Examen JAVA Resuelto - TEST 1a Evaluacion"')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run2 = extra.add_run('\nAula en la nube · MEGA Curso JAVA desde 0')
run2.font.size = Pt(9)
run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECCIÓN: ESTRUCTURA
# ══════════════════════════════════════════════════════════════
add_title('Estructura del Examen', 1)
add_bullet('Parte 1: TEST — 10 preguntas, 30 minutos, SIN apuntes — 2.5 puntos')
add_bullet('Parte 2: Práctico — 2 horas, CON apuntes — 7.5 puntos')
add_bullet('Total: 10 puntos')
add_normal('Único material permitido en el test: tabla ASCII')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# DATOS DEL VÍDEO
# ══════════════════════════════════════════════════════════════
add_title('Datos del Vídeo', 1)
add_bullet('Título: 4-26 DAM - DAW: Examen JAVA Resuelto - TEST 1a Evaluacion')
add_bullet('Canal: Aula en la nube')
add_bullet('Playlist: MEGA Curso JAVA desde 0 [ DAM - DAW ] - TEMA 4')
add_bullet('URL: https://www.youtube.com/watch?v=pUBBuTAd0Co&list=PLG1qdjD__qH6ULjW5iN8E45m5nkaCNbUu&index=92')
add_bullet('Repo: https://github.com/aulaenlanube/curso-programacion-java')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# EJERCICIOS
# ══════════════════════════════════════════════════════════════
# Estructura de cada ejercicio
preguntas = [
    {
        'num': 1,
        'titulo': 'POST-INCREMENTO',
        'codigo': [
            'int a = 1;',
            'a = a++;',
            'Resultado: a = 1'
        ],
        'explicacion': 'a = a++ primero asigna el valor actual (1) y luego incrementa. Como la asignación guarda el valor ANTES de incrementar, a termina siendo 1.'
    },
    {
        'num': 2,
        'titulo': 'CONCATENACIÓN STRING + INT',
        'codigo': [
            'String cad = "AA";',
            'char letra = \'a\';',
            'int num = 1;',
            'cad = cad + letra;     // "AAa"',
            'cad = cad + \'A\';       // "AAaA"',
            'cad = cad + num;       // "AAaA1"',
            'cad += 1 + 2;          // "AAaA13"  (1+2=3, se concatena como String)',
            'Resultado: "AAaA13"'
        ],
        'explicacion': 'Clave: 1+2 se suma primero (3) y luego se concatena al String.'
    },
    {
        'num': 3,
        'titulo': 'BUCLE CON STRING Y LENGTH',
        'codigo': [
            'String cad = "AB";',
            'for (int i = 0; i < 2; i++) {',
            '    cad = cad + 1 + 2;  // sin paréntesis -> concatena "1" y "2"',
            '}',
            'System.out.println(cad + " y tiene " + cad.length() + " caracteres");',
            'Resultado: "AB1212 y tiene 14 caracteres"'
        ],
        'explicacion': '1+2 sin paréntesis en String se concatena como "1"+"2", no se suma. AB + "1" + "2" = AB12 en cada vuelta → AB1212. Longitud: 2+2+2+2+2+2+2 = 14.'
    },
    {
        'num': 4,
        'titulo': 'TABLA ASCII + CARACTERES',
        'codigo': [
            'char letra = 50;        // \'2\' en ASCII',
            'letra *= 2;             // 100 -> \'d\' en ASCII',
            'for (int i = 0; i < 5; i++) {',
            '    letra++;            // 101=\'e\', 102=\'f\', 103=\'g\', 104=\'h\', 105=\'i\'',
            '    System.out.print(letra);',
            '}',
            'System.out.println(": letra final = " + letra + 1);',
            'Resultado: "efghi: letra final = i1"'
        ],
        'explicacion': 'letra+1 al final concatena porque hay un String antes. letra=105 (\'i\'), se concatena con 1 → "i1", no se suma.'
    },
    {
        'num': 5,
        'titulo': 'BARRA r (CARRIAGE RETURN)',
        'codigo': [
            'String cad = "Hola";',
            'for (int i = 0; i < 4; i++) {',
            '    cad += "\\r" + i;     // \\r vuelve al principio de línea',
            '}',
            'Resultado: "3ola"'
        ],
        'explicacion': '\\r mueve el cursor al principio. Cada iteración sobrescribe el primer carácter:\n'
                      '  • Tras i=0: "\\r0" + "Hola" → "0ola"\n'
                      '  • Tras i=1: "\\r1" + "0ola" → "1ola"\n'
                      '  • Tras i=2: "\\r2" + "1ola" → "2ola"\n'
                      '  • Tras i=3: "\\r3" + "2ola" → "3ola"'
    },
    {
        'num': 6,
        'titulo': 'PRE-INCREMENTO + BUCLE',
        'codigo': [
            'int a = 1, b = 2;',
            'if (++a > b++) {',
            '    a++;',
            '}',
            'for (int i = 0; i < a; i++) {',
            '    a = i;',
            '    if (a == i) break;',
            '}',
            'Resultado: a = 0, b = 3'
        ],
        'explicacion': '++a → a=2. b++ → se compara con 2 (b original), luego b=3. 2>2 es false, no entra. '
                      'Luego for: i=0, a=2 → 0<2 true → a=0, como 0==0 → break. a=0, b=3.'
    },
    {
        'num': 7,
        'titulo': 'IFs ANIDADOS + OPERADORES',
        'codigo': [
            'int x = 1, y = 5, z = 10;',
            'if (!false) { }         // siempre true, no hace nada',
            'x += y;                 // x = 6',
            'x--;                    // x = 5',
            'if (x == y) y += x;     // 5==5 -> y = 10',
            'if (!(x != y && y == z)) { }    // !(true && true) = false',
            'if (x == y);            // punto y coma -> if vacío',
            'z -= y;                 // z = 0',
            'x += x;                 // x = 10',
            'Resultado: x=10, y=10, z=0'
        ],
        'explicacion': 'Clave: if(x == y); con punto y coma NO ejecuta la siguiente línea.'
    },
    {
        'num': 8,
        'titulo': 'BUCLES ANIDADOS',
        'codigo': [
            'int a = 0;',
            'for (int i = 1; i < 10; i++)       // 9 vueltas',
            '    for (int j = 1; j < 50; j++)   // 49 vueltas',
            '        a += 2;                     // 9 * 49 * 2 = 882',
            'Resultado: a = 882'
        ],
        'explicacion': '9 × 49 × 2 = 882'
    },
    {
        'num': 9,
        'titulo': 'RECURSIVIDAD',
        'codigo': [
            'static void cadena(String cad, int n) {',
            '    if (n == 3) {',
            '        System.out.println(cad);',
            '    } else {',
            '        cadena(cad + n, n + 1);',
            '    }',
            '}',
            '',
            '// Invocación: cadena("Hola", 0);',
            'Resultado: "Hola012"'
        ],
        'explicacion': 'Traza:\n'
                      '  cadena("Hola", 0) → cadena("Hola0", 1)\n'
                      '  cadena("Hola0", 1) → cadena("Hola01", 2)\n'
                      '  cadena("Hola01", 2) → cadena("Hola012", 3)\n'
                      '  cadena("Hola012", 3) → n==3 → print("Hola012")'
    },
    {
        'num': 10,
        'titulo': 'OPERADOR TERNARIO',
        'codigo': [
            'int a = 1, b = 1;',
            'a = (a != b) ? a + b : a - b;  // a != b? false -> a - b = 0',
            'Resultado: a = 0'
        ],
        'explicacion': 'Como a == b, se evalúa la expresión después de los dos puntos: a - b = 0.'
    },
]

for p in preguntas:
    add_title(f'Pregunta {p["num"]}: {p["titulo"]}', 1)
    add_code('\n'.join(p['codigo']))
    doc.add_paragraph()
    add_normal(p['explicacion'])
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# DISTRIBUCIÓN PUNTUACIÓN
# ══════════════════════════════════════════════════════════════
add_title('Distribución de Puntuación', 1)
add_bullet('Test: 2.5 puntos (25%) — 10 preguntas × 0.25 puntos')
add_bullet('Práctico: 7.5 puntos (75%)')
add_bullet('Total: 10 puntos')

# ── Guardar ──
output = r'C:\Users\Pablo\Desktop\Pablo\anki-cards-completo\pruebas\programacion\JAVA IOC\JAVA YOUTUBE\04-Arrays\Examen_Test_1a_Evaluacion.docx'
doc.save(output)
print(f'Documento creado: {output}')
